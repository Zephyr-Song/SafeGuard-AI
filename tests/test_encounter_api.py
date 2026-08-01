import json
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

from app import app
from modules.encounter_api import encounter_engine
from modules.encounter_engine import EncounterStore, SAMPLE_PROJECT


class UnconfiguredTestLLM:
    providers = {}
    active_provider_id = None

    @staticmethod
    def is_configured():
        return False


class EncounterApiTest(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        encounter_engine.store = EncounterStore(str(Path(self.temp_dir.name) / "api.db"))
        self.original_llm_client = encounter_engine.llm_client
        encounter_engine.llm_client = UnconfiguredTestLLM()
        app.config.update(TESTING=True, SAFEBARS_REQUIRE_ROLE_AUTH=True)
        self.client = app.test_client()

    def tearDown(self):
        encounter_engine.llm_client = self.original_llm_client
        self.temp_dir.cleanup()

    def test_complete_api_workflow(self):
        options = self.client.get("/api/safebars/v2/options")
        self.assertEqual(options.status_code, 200)
        self.assertEqual(len(options.get_json()["scenarios"]), 6)
        self.assertTrue(all(item["trigger_stage"] for item in options.get_json()["scenarios"]))

        payload = json.loads(json.dumps(SAMPLE_PROJECT))
        payload["use_llm"] = False
        payload["artifacts"]["ai_governance"] = "Legacy clients must not add an eleventh public field."
        created = self.client.post("/api/safebars/v2/sessions", json=payload)
        self.assertEqual(created.status_code, 201)
        created_payload = created.get_json()
        self.assertTrue(created_payload["session"]["project"]["uses_ai"])
        self.assertTrue(created_payload["session"]["use_llm"])
        self.assertEqual(created_payload["session"]["artifacts"]["ai_governance"], "")
        session_id = created_payload["session"]["id"]
        researcher_headers = {
            "X-SafeBARS-Access": created_payload["access"]["researcher_token"]
        }
        expert_headers = {
            "X-SafeBARS-Access": created_payload["access"]["expert_token"]
        }

        tradeoffs = self.client.patch(
            f"/api/safebars/v2/sessions/{session_id}/tradeoffs",
            json={
                "deliberations": [
                    {
                        "id": "recruitment_reach",
                        "value": 70,
                        "rationale": "A narrower group is methodologically necessary, with community review of exclusion effects.",
                    },
                    {
                        "id": "data_richness_privacy",
                        "value": 35,
                        "rationale": "Data minimization outweighs richer recordings for this sensitive topic.",
                    },
                ]
            },
            headers=researcher_headers,
        )
        self.assertEqual(tradeoffs.status_code, 200)
        self.assertEqual(
            tradeoffs.get_json()["session"]["tradeoff_deliberations"]["recruitment_reach"]["value"],
            70,
        )

        planned = self.client.patch(
            f"/api/safebars/v2/sessions/{session_id}/plan",
            json={"scenario_ids": ["participant_distress", "partial_withdrawal"]},
            headers=researcher_headers,
        )
        self.assertEqual(planned.status_code, 200)
        planned_session = planned.get_json()["session"]
        self.assertEqual(
            len([task for task in planned_session["audit_plan"] if task["kind"] == "scenario"]),
            2,
        )

        audited = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/audit",
            json={"scenario_ids": ["participant_distress", "partial_withdrawal"]},
            headers=researcher_headers,
        )
        self.assertEqual(audited.status_code, 200)
        audited_session = audited.get_json()["session"]
        self.assertEqual(len(audited_session["traces"]), 2)

        rerunnable = next(
            task
            for task in audited_session["audit_plan"]
            if task.get("scenario_id") == "partial_withdrawal"
        )
        rerun = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/tasks/{rerunnable['id']}/rerun",
            headers=researcher_headers,
        )
        self.assertEqual(rerun.status_code, 200)
        audited_session = rerun.get_json()["session"]

        handoff = audited_session["handoffs"][0]
        handoff_issue = next(
            item for item in audited_session["issues"] if item["id"] == handoff["issue_id"]
        )
        reviewed_passage_id = handoff_issue["source_passage_ids"][0]
        closure_evidence = (
            "A revised withdrawal paragraph names the deletion deadline, process, and responsible contact."
        )
        summary = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/expert-summary",
            headers=expert_headers,
        )
        self.assertEqual(summary.status_code, 200)
        self.assertGreaterEqual(summary.get_json()["summary"]["counts"]["total"], 1)
        expert_review = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/handoffs/{handoff['id']}/review",
            json={
                "action": "advise",
                "reviewer_role": "ethics_board",
                "reviewer_name": "School ethics advisor",
                "advice": "Clarify the withdrawal and deletion boundary.",
                "rationale": "The current procedure is ambiguous.",
                "advice_type": "required_change",
                "responsible_actor": "Principal investigator",
                "closure_evidence": closure_evidence,
                "reviewed_passage_ids": [reviewed_passage_id],
            },
            headers=expert_headers,
        )
        self.assertEqual(expert_review.status_code, 200)
        reviewed_handoff = next(
            item
            for item in expert_review.get_json()["session"]["handoffs"]
            if item["id"] == handoff["id"]
        )
        self.assertEqual(reviewed_handoff["assigned_role"], "ethics_board")
        self.assertEqual(reviewed_handoff["assigned_reviewer_name"], "School ethics advisor")
        self.assertEqual(reviewed_handoff["advice_type"], "required_change")
        self.assertEqual(reviewed_handoff["responsible_actor"], "Principal investigator")
        self.assertEqual(reviewed_handoff["closure_evidence"], closure_evidence)
        self.assertEqual(reviewed_handoff["reviewed_passage_ids"], [reviewed_passage_id])
        self.assertTrue(reviewed_handoff["evidence_reviewed"])
        self.assertGreaterEqual(len(reviewed_handoff["review_history"]), 1)
        reviewed_summary = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/expert-summary",
            headers=expert_headers,
        ).get_json()["summary"]
        self.assertGreaterEqual(reviewed_summary["counts"]["assigned"], 1)
        self.assertIn("ethics_board", reviewed_summary["role_counts"])

        researcher_response = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/handoffs/{handoff['id']}/researcher-response",
            json={
                "response": "We clarified the deletion boundary and escalation owner.",
                "revised_text": "Participants may request deletion until anonymised analysis begins; the study lead records and confirms the request.",
            },
            headers=researcher_headers,
        )
        self.assertEqual(researcher_response.status_code, 200)
        self.assertEqual(
            next(
                item
                for item in researcher_response.get_json()["session"]["handoffs"]
                if item["id"] == handoff["id"]
            )["status"],
            "researcher_revised",
        )
        rereviewed = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/handoffs/{handoff['id']}/review",
            json={
                "action": "advise",
                "reviewer_role": "ethics_board",
                "reviewer_name": "School ethics advisor",
                "advice": "Clarify the withdrawal and deletion boundary.",
                "rationale": "The revised procedure still requires an auditable expert record.",
                "advice_type": "required_change",
                "responsible_actor": "Principal investigator",
                "closure_evidence": closure_evidence,
                "reviewed_passage_ids": [reviewed_passage_id],
            },
            headers=expert_headers,
        )
        self.assertEqual(rereviewed.status_code, 200)
        rereviewed_handoff = next(
            item
            for item in rereviewed.get_json()["session"]["handoffs"]
            if item["id"] == handoff["id"]
        )
        self.assertTrue(rereviewed_handoff["evidence_reviewed"])
        self.assertEqual(
            rereviewed_handoff["reviewed_passage_ids"],
            [reviewed_passage_id],
        )
        protected_rerun = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/audit",
            json={"scenario_ids": ["partial_withdrawal"]},
            headers=researcher_headers,
        )
        self.assertEqual(protected_rerun.status_code, 400)

        versioned = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/versions",
            headers=researcher_headers,
        )
        self.assertEqual(versioned.status_code, 201)
        self.assertEqual(versioned.get_json()["session"]["lineage"]["parent_session_id"], session_id)
        self.assertEqual(versioned.get_json()["session"]["lineage"]["version_number"], 2)
        self.assertTrue(versioned.get_json()["session"]["project"]["uses_ai"])
        self.assertTrue(versioned.get_json()["session"]["use_llm"])
        self.assertEqual(
            versioned.get_json()["session"]["tradeoff_deliberations"]["recruitment_reach"]["value"],
            70,
        )

        issue_id = audited_session["issues"][0]["id"]
        decision = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/issues/{issue_id}/decision",
            json={"decision": "accept", "rationale": "This change is actionable."},
            headers=researcher_headers,
        )
        self.assertEqual(decision.status_code, 200)

        exported = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/export",
            headers=researcher_headers,
        )
        self.assertEqual(exported.status_code, 200)
        self.assertEqual(exported.mimetype, "application/json")
        self.assertEqual(json.loads(exported.data)["id"], session_id)
        self.assertNotIn(created_payload["access"]["researcher_token"], exported.get_data(as_text=True))
        self.assertNotIn(created_payload["access"]["expert_token"], exported.get_data(as_text=True))

        word_report = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/export.docx",
            headers=researcher_headers,
        )
        self.assertEqual(word_report.status_code, 200)
        self.assertEqual(
            word_report.mimetype,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(word_report.data.startswith(b"PK"))
        self.assertIn("full_audit_report.docx", word_report.headers["Content-Disposition"])
        word_document = Document(BytesIO(word_report.data))
        word_text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)
        self.assertIn(SAMPLE_PROJECT["project"]["title"], word_text)
        self.assertIn(SAMPLE_PROJECT["project"]["review_context"], word_text)
        self.assertIn("SAFEBARS FULL AUDIT REPORT", word_text)
        self.assertIn("Framework-grounded ethics map", word_text)
        self.assertIn("Inspectable audit plan", word_text)
        self.assertIn("Resolution rule", word_text)
        self.assertIn("Expert response type", word_text)
        self.assertIn("Required change", word_text)
        self.assertIn("Principal investigator", word_text)
        self.assertIn(closure_evidence, word_text)
        self.assertIn(f"Protocol passages reviewed: {reviewed_passage_id}", word_text)

        pdf_report = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/export.pdf",
            headers=researcher_headers,
        )
        self.assertEqual(pdf_report.status_code, 200)
        self.assertEqual(pdf_report.mimetype, "application/pdf")
        self.assertTrue(pdf_report.data.startswith(b"%PDF-"))
        self.assertGreater(len(pdf_report.data), 5000)
        self.assertIn("full_audit_report.pdf", pdf_report.headers["Content-Disposition"])

        application = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/export.application.docx"
            , headers=researcher_headers
        )
        self.assertEqual(application.status_code, 200)
        application_document = Document(BytesIO(application.data))
        application_text = "\n".join(
            paragraph.text for paragraph in application_document.paragraphs
        )
        self.assertIn("ETHICS APPLICATION DRAFT", application_text)
        self.assertIn("formal approval", application_text)
        self.assertIn("Research area and ethics-review context", application_text)
        self.assertIn(SAMPLE_PROJECT["project"]["review_context"], application_text)
        self.assertIn("Required change", application_text)
        self.assertIn("Principal investigator", application_text)
        self.assertIn(closure_evidence, application_text)
        self.assertIn(
            f"Protocol passages reviewed: {reviewed_passage_id}",
            application_text,
        )

        research_design = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/export.research-design.docx",
            headers=researcher_headers,
        )
        self.assertEqual(research_design.status_code, 200)
        self.assertIn("research_design.docx", research_design.headers["Content-Disposition"])
        design_document = Document(BytesIO(research_design.data))
        design_text = "\n".join(paragraph.text for paragraph in design_document.paragraphs)
        self.assertIn("RESEARCH DESIGN AND ETHICS-IN-PRACTICE PLAN", design_text)
        self.assertIn(SAMPLE_PROJECT["project"]["review_context"], design_text)
        self.assertIn("Research setting, procedures, and participant journey", design_text)
        self.assertIn("Ethics-informed trade-offs and design decisions", design_text)
        self.assertIn("Expert dependencies and unresolved design questions", design_text)
        self.assertIn("Narrow eligibility: 70", design_text)
        self.assertIn("A narrower group is methodologically necessary", design_text)
        self.assertIn("Required change", design_text)
        self.assertIn("Principal investigator", design_text)
        self.assertIn(closure_evidence, design_text)
        self.assertIn(f"Protocol passages reviewed: {reviewed_passage_id}", design_text)

        expert_export = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/export.expert.docx"
            , headers=expert_headers
        )
        self.assertEqual(expert_export.status_code, 200)
        expert_document = Document(BytesIO(expert_export.data))
        expert_text = "\n".join(paragraph.text for paragraph in expert_document.paragraphs)
        self.assertIn("SAFEBARS EXPERT REVIEW SUMMARY", expert_text)
        self.assertIn("Expert advice", expert_text)
        self.assertIn("Expert response type", expert_text)
        self.assertIn("Required change", expert_text)
        self.assertIn("Principal investigator", expert_text)
        self.assertIn(closure_evidence, expert_text)
        self.assertIn(f"Protocol passages reviewed: {reviewed_passage_id}", expert_text)

        second_payload = json.loads(json.dumps(SAMPLE_PROJECT))
        second_payload["project"]["title"] = "Second sensitive-service protocol"
        second_created = self.client.post(
            "/api/safebars/v2/sessions", json=second_payload
        ).get_json()
        portfolio = self.client.post(
            "/api/safebars/v2/expert/export.portfolio.docx",
            json={
                "cases": [
                    {
                        "session_id": session_id,
                        "expert_token": created_payload["access"]["expert_token"],
                    },
                    {
                        "session_id": second_created["session"]["id"],
                        "expert_token": second_created["access"]["expert_token"],
                    },
                ]
            },
        )
        self.assertEqual(portfolio.status_code, 200)
        self.assertIn("expert_caseload_summary.docx", portfolio.headers["Content-Disposition"])
        portfolio_document = Document(BytesIO(portfolio.data))
        portfolio_text = "\n".join(
            paragraph.text for paragraph in portfolio_document.paragraphs
        )
        self.assertIn("SAFEBARS EXPERT CASELOAD SUMMARY", portfolio_text)
        self.assertIn("Applications in scope: 2", portfolio_text)
        self.assertIn(SAMPLE_PROJECT["project"]["title"], portfolio_text)
        self.assertIn("Second sensitive-service protocol", portfolio_text)
        self.assertIn("Clarify the withdrawal and deletion boundary.", portfolio_text)
        self.assertIn("Required change", portfolio_text)
        self.assertIn("Principal investigator", portfolio_text)
        self.assertIn(closure_evidence, portfolio_text)
        self.assertIn(f"Protocol passages reviewed: {reviewed_passage_id}", portfolio_text)
        self.assertIn("Recorded research-design trade-offs", portfolio_text)
        self.assertIn("Data minimization outweighs richer recordings", portfolio_text)

    def test_structured_expert_review_api_enforces_closure_contract(self):
        payload = json.loads(json.dumps(SAMPLE_PROJECT))
        payload["use_llm"] = False
        created = self.client.post("/api/safebars/v2/sessions", json=payload)
        self.assertEqual(created.status_code, 201)
        created_payload = created.get_json()
        session_id = created_payload["session"]["id"]
        researcher_headers = {
            "X-SafeBARS-Access": created_payload["access"]["researcher_token"]
        }
        expert_headers = {
            "X-SafeBARS-Access": created_payload["access"]["expert_token"]
        }

        audited = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/audit",
            json={"scenario_ids": ["partial_withdrawal"]},
            headers=researcher_headers,
        )
        self.assertEqual(audited.status_code, 200)
        audited_session = audited.get_json()["session"]
        handoff = next(
            item
            for item in audited_session["handoffs"]
            if item["issue_id"] == "issue_partial_withdrawal"
        )
        issue = next(
            item for item in audited_session["issues"] if item["id"] == handoff["issue_id"]
        )
        reviewed_passage_id = issue["source_passage_ids"][0]
        review_payload = {
            "reviewer_role": "ethics_board",
            "reviewer_name": "School ethics advisor",
            "advice_type": "required_change",
            "advice": "Revise the withdrawal procedure before recruitment begins.",
            "rationale": "Participants need a clear and usable boundary for deletion requests.",
            "responsible_actor": "Principal investigator",
            "closure_evidence": (
                "A revised protocol paragraph states the deadline, contact, and deletion process."
            ),
            "reviewed_passage_ids": [reviewed_passage_id],
        }
        review_url = (
            f"/api/safebars/v2/sessions/{session_id}/handoffs/{handoff['id']}/review"
        )

        missing_field = self.client.post(
            review_url,
            json={**review_payload, "action": "advise", "closure_evidence": ""},
            headers=expert_headers,
        )
        self.assertEqual(missing_field.status_code, 400)
        self.assertIn("evidence would be sufficient", missing_field.get_json()["error"])

        forged_evidence = self.client.post(
            review_url,
            json={
                **review_payload,
                "action": "advise",
                "reviewed_passage_ids": ["CON-999"],
            },
            headers=expert_headers,
        )
        self.assertEqual(forged_evidence.status_code, 400)
        self.assertIn("unknown protocol passage", forged_evidence.get_json()["error"])

        after_rejections = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/expert-summary",
            headers=expert_headers,
        )
        rejected_handoff = next(
            item
            for item in after_rejections.get_json()["summary"]["queue"]
            if item["id"] == handoff["id"]
        )
        self.assertEqual(rejected_handoff["status"], "open")
        self.assertEqual(rejected_handoff["review_history"], [])

        advised = self.client.post(
            review_url,
            json={**review_payload, "action": "advise"},
            headers=expert_headers,
        )
        self.assertEqual(advised.status_code, 200)

        premature_close = self.client.post(
            review_url,
            json={**review_payload, "action": "resolve"},
            headers=expert_headers,
        )
        self.assertEqual(premature_close.status_code, 400)
        self.assertIn("has not responded", premature_close.get_json()["error"])

        researcher_response = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/handoffs/{handoff['id']}/researcher-response",
            json={
                "response": (
                    "The deletion deadline, responsible contact, and confirmation step are now explicit."
                ),
                "revised_text": (
                    "Participants may request deletion until anonymised analysis begins by "
                    "contacting the principal investigator."
                ),
            },
            headers=researcher_headers,
        )
        self.assertEqual(researcher_response.status_code, 200)
        revised_handoff = next(
            item
            for item in researcher_response.get_json()["session"]["handoffs"]
            if item["id"] == handoff["id"]
        )
        self.assertEqual(revised_handoff["status"], "researcher_revised")
        self.assertEqual(revised_handoff["reviewed_passage_ids"], [])
        self.assertFalse(revised_handoff["evidence_reviewed"])
        self.assertFalse(revised_handoff["evidence_reviewed_at"])

        stale_summary = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/expert-summary",
            headers=expert_headers,
        ).get_json()["summary"]
        stale_handoff = next(
            item for item in stale_summary["queue"] if item["id"] == handoff["id"]
        )
        self.assertFalse(stale_handoff["evidence_review_current"])
        self.assertFalse(stale_handoff["closure_record_complete"])

        close_without_new_evidence_review = self.client.post(
            review_url,
            json={
                "action": "resolve",
                "reviewer_role": review_payload["reviewer_role"],
                "reviewer_name": review_payload["reviewer_name"],
            },
            headers=expert_headers,
        )
        self.assertEqual(close_without_new_evidence_review.status_code, 400)
        self.assertIn(
            "review at least one cited protocol passage",
            close_without_new_evidence_review.get_json()["error"],
        )

        resolved = self.client.post(
            review_url,
            json={**review_payload, "action": "resolve"},
            headers=expert_headers,
        )
        self.assertEqual(resolved.status_code, 200)
        resolved_handoff = next(
            item
            for item in resolved.get_json()["session"]["handoffs"]
            if item["id"] == handoff["id"]
        )
        self.assertEqual(resolved_handoff["status"], "resolved")

        summary = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/expert-summary",
            headers=expert_headers,
        ).get_json()["summary"]
        closed_summary = next(
            item for item in summary["queue"] if item["id"] == handoff["id"]
        )
        self.assertTrue(closed_summary["closure_record_complete"])
        self.assertFalse(closed_summary["legacy_resolution"])
        self.assertEqual(summary["counts"]["resolved"], 1)
        self.assertEqual(summary["counts"]["unresolved"], 0)

        state_before_rejected_response = {
            "status": closed_summary["status"],
            "resolved_at": closed_summary["resolved_at"],
            "researcher_response": closed_summary.get("researcher_response", ""),
            "researcher_revised_text": closed_summary.get(
                "researcher_revised_text",
                "",
            ),
            "researcher_revision_history": list(
                closed_summary.get("researcher_revision_history", [])
            ),
            "review_history": list(closed_summary["review_history"]),
            "reviewed_passage_ids": list(closed_summary["reviewed_passage_ids"]),
            "evidence_reviewed": closed_summary["evidence_reviewed"],
            "closure_record_complete": closed_summary["closure_record_complete"],
        }
        rejected_response = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/handoffs/{handoff['id']}/researcher-response",
            json={
                "response": "A later response must not silently reopen a closed handoff.",
                "revised_text": "This text must not be stored.",
            },
            headers=researcher_headers,
        )
        self.assertEqual(rejected_response.status_code, 400)
        self.assertIn("handoff is closed", rejected_response.get_json()["error"])

        unchanged_summary = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/expert-summary",
            headers=expert_headers,
        ).get_json()["summary"]
        unchanged_handoff = next(
            item
            for item in unchanged_summary["queue"]
            if item["id"] == handoff["id"]
        )
        state_after_rejected_response = {
            "status": unchanged_handoff["status"],
            "resolved_at": unchanged_handoff["resolved_at"],
            "researcher_response": unchanged_handoff.get("researcher_response", ""),
            "researcher_revised_text": unchanged_handoff.get(
                "researcher_revised_text",
                "",
            ),
            "researcher_revision_history": list(
                unchanged_handoff.get("researcher_revision_history", [])
            ),
            "review_history": list(unchanged_handoff["review_history"]),
            "reviewed_passage_ids": list(unchanged_handoff["reviewed_passage_ids"]),
            "evidence_reviewed": unchanged_handoff["evidence_reviewed"],
            "closure_record_complete": unchanged_handoff[
                "closure_record_complete"
            ],
        }
        self.assertEqual(
            state_after_rejected_response,
            state_before_rejected_response,
        )

    def test_legacy_resolved_handoff_summary_and_export_remain_compatible(self):
        payload = json.loads(json.dumps(SAMPLE_PROJECT))
        payload["use_llm"] = False
        created = self.client.post("/api/safebars/v2/sessions", json=payload)
        self.assertEqual(created.status_code, 201)
        created_payload = created.get_json()
        session_id = created_payload["session"]["id"]
        researcher_headers = {
            "X-SafeBARS-Access": created_payload["access"]["researcher_token"]
        }
        expert_headers = {
            "X-SafeBARS-Access": created_payload["access"]["expert_token"]
        }
        audited = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/audit",
            json={"scenario_ids": ["partial_withdrawal"]},
            headers=researcher_headers,
        )
        self.assertEqual(audited.status_code, 200)
        legacy_session = audited.get_json()["session"]
        legacy_handoff = legacy_session["handoffs"][0]
        legacy_handoff["status"] = "resolved"
        legacy_handoff["resolved_at"] = "2025-01-01T00:00:00+00:00"
        legacy_handoff["expert_advice"] = (
            "Legacy advice recorded before structured closure fields existed."
        )
        legacy_handoff["expert_rationale"] = (
            "Legacy ethical rationale retained for the audit trail."
        )
        legacy_handoff["review_history"] = [
            {
                "action": "resolve",
                "reviewer_role": "ethics_board",
                "reviewer_name": "Legacy reviewer",
                "advice": legacy_handoff["expert_advice"],
                "rationale": legacy_handoff["expert_rationale"],
                "timestamp": "2025-01-01T00:00:00+00:00",
            }
        ]
        for field in (
            "advice_type",
            "advice_type_label",
            "responsible_actor",
            "closure_evidence",
            "reviewed_passage_ids",
            "evidence_gap_acknowledged",
            "evidence_reviewed",
            "evidence_reviewed_at",
        ):
            legacy_handoff.pop(field, None)
        encounter_engine.store.save(legacy_session)

        summary_response = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/expert-summary",
            headers=expert_headers,
        )
        self.assertEqual(summary_response.status_code, 200)
        summary = summary_response.get_json()["summary"]
        legacy_summary = next(
            item for item in summary["queue"] if item["id"] == legacy_handoff["id"]
        )
        self.assertEqual(legacy_summary["status"], "resolved")
        self.assertFalse(legacy_summary["closure_record_complete"])
        self.assertTrue(legacy_summary["legacy_resolution"])

        expert_export = self.client.get(
            f"/api/safebars/v2/sessions/{session_id}/export.expert.docx",
            headers=expert_headers,
        )
        self.assertEqual(expert_export.status_code, 200)
        document = Document(BytesIO(expert_export.data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Legacy advice recorded", text)
        self.assertIn("No reviewed protocol evidence was recorded.", text)

    def test_role_tokens_and_expert_invite_rotation(self):
        payload = json.loads(json.dumps(SAMPLE_PROJECT))
        created = self.client.post("/api/safebars/v2/sessions", json=payload).get_json()
        session_id = created["session"]["id"]
        researcher_token = created["access"]["researcher_token"]
        expert_token = created["access"]["expert_token"]
        researcher_headers = {"X-SafeBARS-Access": researcher_token}
        expert_headers = {"X-SafeBARS-Access": expert_token}

        self.assertEqual(
            self.client.get(f"/api/safebars/v2/sessions/{session_id}").status_code,
            401,
        )
        self.assertEqual(
            self.client.get(
                f"/api/safebars/v2/sessions/{session_id}", headers=expert_headers
            ).status_code,
            403,
        )
        self.assertEqual(
            self.client.get(
                f"/api/safebars/v2/sessions/{session_id}/expert-summary",
                headers=researcher_headers,
            ).status_code,
            403,
        )
        self.assertEqual(
            self.client.post(
                "/api/safebars/v2/expert/export.portfolio.docx",
                json={
                    "cases": [
                        {
                            "session_id": session_id,
                            "expert_token": researcher_token,
                        }
                    ]
                },
            ).status_code,
            403,
        )

        rotated = self.client.post(
            f"/api/safebars/v2/sessions/{session_id}/access/rotate-expert",
            headers=researcher_headers,
        )
        self.assertEqual(rotated.status_code, 200)
        new_expert_headers = {
            "X-SafeBARS-Access": rotated.get_json()["expert_token"]
        }
        self.assertEqual(
            self.client.get(
                f"/api/safebars/v2/sessions/{session_id}/expert-summary",
                headers=expert_headers,
            ).status_code,
            403,
        )
        self.assertEqual(
            self.client.get(
                f"/api/safebars/v2/sessions/{session_id}/expert-summary",
                headers=new_expert_headers,
            ).status_code,
            200,
        )

    def test_workspace_and_v1_routes_are_available(self):
        workspace = self.client.get("/safebars")
        self.assertEqual(workspace.status_code, 200)
        self.assertIn(
            "https://unpkg.com",
            workspace.headers.get("Content-Security-Policy", ""),
        )
        self.assertIn(b"Encounter stress-testing workspace", workspace.data)
        self.assertIn(b"function refreshIcons()", workspace.data)
        self.assertIn(b'integrity="sha384-', workspace.data)
        self.assertIn(b"Start quick intake", workspace.data)
        self.assertIn(b"How to use SafeBARS", workspace.data)
        self.assertIn(b'aria-label="Close instructions"', workspace.data)
        self.assertIn(b'class="dialog-close-symbol"', workspace.data)
        self.assertIn(b"For researchers", workspace.data)
        self.assertIn(b"For ethics experts", workspace.data)
        self.assertIn(b"Keep your record", workspace.data)
        self.assertIn(b"Six short questions", workspace.data)
        self.assertIn(b"Research area and ethics-review context", workspace.data)
        self.assertIn(
            b'<textarea id="projectTitle" class="project-scroll-field" rows="3"',
            workspace.data,
        )
        self.assertIn(
            b'<textarea id="projectReviewContext" class="project-scroll-field" rows="3"',
            workspace.data,
        )
        self.assertIn(
            b'<textarea id="projectContext" class="project-scroll-field" rows="3"',
            workspace.data,
        )
        self.assertIn(
            b'<textarea id="targetPeople" class="project-scroll-field" rows="3"',
            workspace.data,
        )
        self.assertIn(b"height: 72px", workspace.data)
        self.assertIn(b"min-height: 72px", workspace.data)
        self.assertIn(b"overflow-y: scroll", workspace.data)
        self.assertIn(b"resize: vertical", workspace.data)
        self.assertIn(b"scrollbar-gutter: stable", workspace.data)
        self.assertIn(b"scrollbar-color:", workspace.data)
        self.assertIn(b".project-scroll-field::-webkit-scrollbar-thumb", workspace.data)
        self.assertIn(
            b'use the mouse wheel or drag the scrollbar on the right',
            workspace.data,
        )
        self.assertEqual(
            workspace.data.count(b'aria-describedby="projectFieldScrollHelp"'),
            4,
        )
        self.assertNotIn(b"max-height: 72px", workspace.data)
        self.assertNotIn(b"What are you studying", workspace.data)
        self.assertIn(b"Intake completed", workspace.data)
        self.assertIn(b"Done \xc2\xb7 review populated fields", workspace.data)
        self.assertIn(b"Full audit report (.docx)", workspace.data)
        self.assertIn(b"Application draft (.docx)", workspace.data)
        self.assertIn(b"Research design (.docx)", workspace.data)
        self.assertIn(b"Expert workspace", workspace.data)
        self.assertIn(b'id="expertWorkspaceButton"', workspace.data)
        self.assertIn(b"Open expert review", workspace.data)
        self.assertNotIn(b"no handoffs yet", workspace.data)
        self.assertIn(b"Plain-language guide to interface terms", workspace.data)
        self.assertIn(b"Internal source ID", workspace.data)
        self.assertIn(b"Evidence from", workspace.data)
        self.assertIn(b"How to read a source label", workspace.data)
        self.assertIn(b"Needs attention", workspace.data)
        self.assertIn(b"Closest submitted material to inspect", workspace.data)
        self.assertIn(b"Show in protocol", workspace.data)
        self.assertIn(b"Connected design decisions", workspace.data)
        self.assertIn(b'id="readinessCheckButton"', workspace.data)
        self.assertIn(b'id="readinessCheckDialog"', workspace.data)
        self.assertIn(b"Review readiness", workspace.data)
        self.assertIn(b"See blockers and the next action", workspace.data)
        self.assertIn(b"function readinessSnapshot()", workspace.data)
        self.assertIn(b"function followReadinessNextAction()", workspace.data)
        self.assertNotIn(b'href="/safebars/v1"', workspace.data)
        self.assertNotIn(b"Research rehearsal", workspace.data)
        self.assertIn(b"frameworkCoverageSummary", workspace.data)
        self.assertIn(b"frameworkCoverageDonut", workspace.data)
        self.assertIn(b"Evidence coverage at a glance", workspace.data)
        self.assertIn(b"Participant journey stress map", workspace.data)
        self.assertIn(b'data-journey-lens="combined"', workspace.data)
        self.assertIn(b"Open scenario result", workspace.data)
        self.assertIn(b"Unsaved material changes", workspace.data)
        self.assertIn(b"Excluded from scope", workspace.data)
        self.assertIn(b"hadUnsavedMaterialChanges", workspace.data)
        self.assertNotIn(b"Project researches or uses AI", workspace.data)
        self.assertNotIn(b"Optional bounded LLM critic", workspace.data)
        self.assertNotIn(b'id="usesAi"', workspace.data)
        self.assertNotIn(b'id="useLlm"', workspace.data)
        self.assertNotIn(b'id="llmLabel"', workspace.data)
        self.assertIn(b"AI review \xc2\xb7 automatic", workspace.data)
        self.assertNotIn(b"Zhipu", workspace.data)
        self.assertIn(b'id="artifactEditorSelect"', workspace.data)
        self.assertIn(b"0 of 10 fields", workspace.data)
        self.assertIn(b"1 of 6 \xc2\xb7 not added", workspace.data)
        self.assertIn(b"keeps all six connected", workspace.data)
        self.assertNotIn(b'data-artifact-panel="ai_governance"', workspace.data)
        self.assertNotIn(b'id="artifactAiGovernance"', workspace.data)
        self.assertNotIn(b"AI ethics-review supplement", workspace.data)
        self.assertIn(b'for="artifactRecruitment"', workspace.data)
        self.assertIn(b"Run this check again", workspace.data)
        self.assertIn(b"The plan has already run", workspace.data)
        self.assertNotIn(b">V1</span>", workspace.data)
        self.assertNotIn(b">Rerun</button>", workspace.data)
        self.assertIn(b"window.location.assign(expertInviteUrl())", workspace.data)
        self.assertIn(b"Framework-linked Trade-off Board", workspace.data)
        self.assertIn(b"Downstream outputs", workspace.data)
        self.assertIn(b"10.1145/3334480.3382795", workspace.data)

        legacy = self.client.get("/safebars/v1")
        self.assertEqual(legacy.status_code, 200)
        self.assertIn(b"Rehearsal Chat", legacy.data)

        expert = self.client.get("/safebars/expert/example_session")
        self.assertEqual(expert.status_code, 200)
        self.assertIn(b"SafeBARS Expert Workspace", expert.data)
        self.assertIn(b"All expert roles", expert.data)
        self.assertIn(b"Workflow timeline and comment history", expert.data)
        self.assertIn(b"Needs expert action", expert.data)
        self.assertIn(b"The decision you are being asked to make", expert.data)
        self.assertIn(b"Downstream outputs", expert.data)
        self.assertIn(b"Enter one auditable expert response", expert.data)
        self.assertIn(b"Expert response type", expert.data)
        self.assertIn(b"Responsible person or role", expert.data)
        self.assertIn(b"Evidence sufficient to close", expert.data)
        self.assertIn(b"Closure requirements", expert.data)
        self.assertIn(b"these are not checkboxes", expert.data)
        self.assertIn(b"Decision-support boundary", expert.data)
        self.assertIn(b"function decisionReadiness", expert.data)
        self.assertIn(b"function markEvidenceReviewed", expert.data)
        self.assertIn(b"Record evidence review", expert.data)
        self.assertIn(b"data-record-evidence", expert.data)
        self.assertNotIn(b"Evidence review is recorded automatically", expert.data)
        self.assertIn(
            b'canAsk: item.status !== "resolved" && coreComplete && values.advice_type === "clarification_request"',
            expert.data,
        )
        self.assertIn(b"function isWaiting(item)", expert.data)
        self.assertIn(b"const responseAt = latestResearcherEvidenceAt(item)", expert.data)
        self.assertIn(b"const requestAt = latestBlockingRequestAt(item)", expert.data)
        self.assertIn(
            b"timestampValue(responseAt) <= timestampValue(requestAt)",
            expert.data,
        )
        self.assertIn(b"function defaultTabFor(item)", expert.data)
        self.assertIn(
            b'if (item.status === "resolved" || isWaiting(item)) return "history"',
            expert.data,
        )
        self.assertIn(
            b'if (["researcher_revised","researcher_responded"].includes(item.status)) return "evidence"',
            expert.data,
        )
        self.assertIn(b"activeTab = defaultTabFor(queue[0])", expert.data)
        self.assertIn(b"activeTab = defaultTabFor(next)", expert.data)
        expert_html = expert.get_data(as_text=True)
        review_function_start = expert_html.index("async function review(id, action)")
        body_start = expert_html.index("const body = {", review_function_start)
        substantive_guard = expert_html.index(
            'if (["advise","request_clarification","resolve"].includes(action))',
            body_start,
        )
        redirect_payload = expert_html.index(
            'if (action === "redirect") body.redirect_role',
            substantive_guard,
        )
        base_body = expert_html[body_start:substantive_guard]
        substantive_body = expert_html[substantive_guard:redirect_payload]
        for structured_field in (
            "advice_type",
            "advice:",
            "rationale:",
            "responsible_actor",
            "closure_evidence",
            "reviewed_passage_ids",
            "evidence_gap_acknowledged",
        ):
            self.assertNotIn(structured_field, base_body)
            self.assertIn(structured_field, substantive_body)
        reopen_start = expert_html.index(
            'if (action === "reopen") {',
            redirect_payload,
        )
        reopen_end = expert_html.index("await loadSummary()", reopen_start)
        reopen_block = expert_html[reopen_start:reopen_end]
        self.assertIn(
            "sessionStorage.removeItem(evidenceReviewKey(id))",
            reopen_block,
        )
        self.assertNotIn(b"Before you save a decision", expert.data)
        self.assertNotIn(
            b"I checked the cited protocol evidence, not only the AI summary.",
            expert.data,
        )
        self.assertIn(b"Expert response \xc2\xb7 enter advice", expert.data)
        self.assertIn(b"function openDecisionForm", expert.data)
        self.assertIn(b"field.scrollIntoView", expert.data)
        self.assertIn(b"Send advice to researcher", expert.data)
        self.assertIn(b"Ask researcher", expert.data)
        self.assertIn(b"Redirect handoff", expert.data)
        self.assertIn(b"Close handoff", expert.data)
        self.assertIn(b"Advice sent to the researcher", expert.data)
        self.assertIn(b"Internal source ID", expert.data)
        self.assertIn(b"source ID", expert.data)

        expert_dashboard = self.client.get("/safebars/expert")
        self.assertEqual(expert_dashboard.status_code, 200)
        self.assertIn(b"SafeBARS Expert Caseload", expert_dashboard.data)
        self.assertIn(b"Download caseload summary", expert_dashboard.data)
        self.assertIn(b"filter by assigned or recommended expert role", expert_dashboard.data)
        self.assertIn(b"No review case is open", expert_dashboard.data)
        self.assertIn(b"not an empty input form", expert_dashboard.data)


    def test_tradeoffs_rejects_out_of_range_and_unknown_ids(self):
        payload = json.loads(json.dumps(SAMPLE_PROJECT))
        payload["use_llm"] = False
        created = self.client.post("/api/safebars/v2/sessions", json=payload)
        self.assertEqual(created.status_code, 201)
        researcher_headers = {
            "X-SafeBARS-Access": created.get_json()["access"]["researcher_token"]
        }
        session_id = created.get_json()["session"]["id"]

        # Out-of-range position -> 400 (server-side validation, never trusted blindly).
        bad_range = self.client.patch(
            f"/api/safebars/v2/sessions/{session_id}/tradeoffs",
            json={"deliberations": [{"id": "recruitment_reach", "value": 250}]},
            headers=researcher_headers,
        )
        self.assertEqual(bad_range.status_code, 400)

        # Only unrecognized ids -> 400.
        bad_ids = self.client.patch(
            f"/api/safebars/v2/sessions/{session_id}/tradeoffs",
            json={"deliberations": [{"id": "ghost", "value": 50}]},
            headers=researcher_headers,
        )
        self.assertEqual(bad_ids.status_code, 400)

        # A mixed payload with one valid id still saves the valid one (200).
        mixed = self.client.patch(
            f"/api/safebars/v2/sessions/{session_id}/tradeoffs",
            json={
                "deliberations": [
                    {"id": "recruitment_reach", "value": 65},
                    {"id": "ghost", "value": 50},
                ]
            },
            headers=researcher_headers,
        )
        self.assertEqual(mixed.status_code, 200)
        self.assertIn(
            "recruitment_reach",
            mixed.get_json()["session"]["tradeoff_deliberations"],
        )


if __name__ == "__main__":
    unittest.main()
