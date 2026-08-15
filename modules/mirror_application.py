"""Build a university-ethics-committee application draft from a Mirror session.

The Ethical Mirror helps a researcher anticipate unintended consequences of an
AI / data-driven research plan *before* implementation.  This module turns the
reflection that the Mirror captures (research plan, literature-grounded lenses,
role probes, tension map, researcher revisions) into a structured draft that
the researcher can hand to an institutional ethics committee.

It deliberately follows the Ethics and Society Review (ESR; Bernstein et al.,
PNAS 2021) framing: standard review (for example an IRB under the Common Rule)
focuses on risks to individual participants, but research can also create risks
to society, to subgroups, and globally.  The draft therefore surfaces those
societal / subgroup / global risks and asks the researcher to commit to
mitigations.  The document is a DRAFT and never an approval or compliance
verdict.
"""

from __future__ import annotations

import io
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
AMBER = RGBColor(0x7A, 0x5A, 0x00)

INTAKE_LABELS: Dict[str, str] = {
    "research_context": "Research context and purpose",
    "intended_change": "Intended benefit / change",
    "direct_users": "Direct users",
    "ai_role": "Role of AI in the research",
    "data_materials": "Data and materials",
    "sensitive_data_justification": "Sensitive data justification",
    "affected_others": "Indirectly affected others",
    "value_commitment": "Value commitment",
    "stop_condition": "Redesign / stop threshold",
    "optional_perspective_context": "Additional perspective context",
}

# Lens statuses ordered from weakest to strongest evidence of reflection.
WEAK_LENS_STATUSES = {"missing", "claimed"}


def _text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, (list, tuple)):
        return "; ".join(str(v) for v in value if v)
    if isinstance(value, dict):
        return "; ".join(str(v) for v in value.values() if v)
    return str(value)


def _set_run(run, size=None, color=None, bold=False, italic=False):
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _label_paragraph(document: Document, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    _set_run(paragraph.add_run(f"{label}: "), bold=True)
    paragraph.add_run(_text(value))
    return None


def _add_esr_section(document: Document, session: Dict[str, Any]) -> None:
    """ESR-style societal / subgroup / global risk statement (Bernstein 2021)."""

    document.add_heading(
        "5. Societal & community risk statement (Ethics and Society Review)",
        level=1,
    )
    paragraph = document.add_paragraph()
    _set_run(paragraph.add_run(
        "Standard ethics review (for example an IRB under the US Common Rule) "
        "focuses on risks to individual participants. The Ethics and Society "
        "Review (ESR; Bernstein, Levi, Magnus, Rajala, Satz & Waeiss, PNAS 2021, "
        "DOI 10.1073/pnas.2117261118) argues that research can also create risks "
        "to society, to subgroups within society, and globally, and that "
        "researchers should state those risks and commit to mitigation "
        "strategies before the work proceeds."
    ))
    document.add_paragraph(
        "Use the prompts below to state, in your own words, who might be "
        "affected beyond the direct participants, what could go wrong, and how "
        "you will mitigate it. Leave no blank line unaddressed before submission."
    )

    # Surface the people/roles the researcher already named as prompts.
    prompts: List[str] = []
    intake = session.get("intake_answers") or {}
    if _text(intake.get("direct_users")):
        prompts.append(
            f"Direct users you named: {_text(intake.get('direct_users'))}. "
            "Who else could be affected, and how?"
        )
    if _text(intake.get("affected_others")):
        prompts.append(
            f"Indirectly affected others you named: "
            f"{_text(intake.get('affected_others'))}. "
            "What subgroup or community-level harms could arise?"
        )
    if not prompts:
        prompts.append(
            "Who beyond the direct participants could be affected by this "
            "research, at the subgroup or societal level?"
        )
    prompts.append(
        "What global, long-term, or systemic consequence could this work "
        "produce, even if individual participants are protected?"
    )
    prompts.append(
        "Which literature lens below is 'Missing' or only 'Claimed' (not yet "
        "reasoned or action-linked)? State the risk it points to and your "
        "mitigation commitment."
    )

    for index, prompt in enumerate(prompts, start=1):
        item = document.add_paragraph(style="List Number")
        _set_run(item.add_run(prompt))
        blank = document.add_paragraph()
        _set_run(blank.add_run("Mitigation statement (researcher to complete): "), bold=True)
        blank.add_run("_" * 60)

    # Echo the weak lenses so the researcher sees exactly which dimensions are open.
    lenses = session.get("lenses") or []
    weak = [l for l in lenses if l.get("status") in WEAK_LENS_STATUSES]
    if weak:
        note = document.add_paragraph()
        _set_run(note.add_run("Open ethical dimensions flagged by the Mirror:"), bold=True)
        for lens in weak:
            line = document.add_paragraph(style="List Bullet")
            _set_run(line.add_run(
                f"{_text(lens.get('label') or lens.get('id'), 'Lens')} "
                f"[{_text(lens.get('status'))}] — {_text(lens.get('evidence') or lens.get('interpretation'))}"
            ))


def _add_deep_audit_section(document: Document, session: Dict[str, Any]) -> None:
    """Surface the deep-audit discovery cues in the committee draft.

    These cues are hypothesis-generating: they surface risks the plan did not
    state, using real failure cases, curated patterns, and signal checks. They
    are not findings or approval.
    """
    document.add_heading("4. Deep-audit discovery cues", level=1)
    da = session.get("deep_audit") or {}
    if not da or not da.get("counts"):
        document.add_paragraph(
            "No deep-audit cues were generated for this plan."
        )
        return
    document.add_paragraph(
        "The Ethical Mirror's deep-audit layer surfaces blind spots beyond what "
        "the submitted plan states. These are discovery cues, not ethics "
        "findings. Each still requires human judgement and real-stakeholder "
        "verification before submission."
    )

    domains = (da.get("domain_flags") or {}).get("matched_domains") or []
    if domains:
        document.add_paragraph().add_run(
            "Detected high-risk domains (with auto-injected checklist):"
        ).bold = True
        for d in domains:
            item = document.add_paragraph(style="List Bullet")
            _set_run(item.add_run(f"{_text(d.get('label'))} [{_text(d.get('severity'))}]: "), bold=True)
            item.add_run(_text(d.get("matched_terms")))
        for c in (da.get("domain_flags") or {}).get("injected_checklist") or []:
            item = document.add_paragraph(style="List Bullet 2")
            _set_run(item.add_run(f"{_text(c.get('title'))} ({_text(c.get('regulation'))}): "), bold=True)
            item.add_run(_text(c.get("detail")))

    ctr = da.get("internal_contradictions") or []
    if ctr:
        document.add_paragraph().add_run("Internal contradictions in the plan:").bold = True
        for c in ctr:
            item = document.add_paragraph(style="List Bullet")
            _set_run(item.add_run(f"{_text(c.get('id'))} — {_text(c.get('explanation'))} "), bold=True)
            item.add_run(f"Commitment: “{_text(c.get('commitment_text'))}”. Fix: {_text(c.get('suggested_action'))}")

    gaps = (da.get("severity_weighted_gaps") or {}).get("prioritized_gaps") or []
    if gaps:
        document.add_paragraph().add_run("Severity-weighted coverage gaps (prioritise):").bold = True
        for g in gaps:
            item = document.add_paragraph(style="List Bullet")
            _set_run(item.add_run(f"{_text(g.get('label'))} [{_text(g.get('severity'))}]"), bold=True)
            if g.get("reasons"):
                item.add_run(f" — {_text(g.get('reasons'))}")

    edges = da.get("additional_dissonance_edges") or []
    if edges:
        document.add_paragraph().add_run("Discovery tensions:").bold = True
        for e in edges:
            item = document.add_paragraph(style="List Bullet")
            _set_run(item.add_run(f"{_text(e.get('id'))} [{_text(e.get('rule'))}]: "), bold=True)
            item.add_run(_text(e.get("tension")))

    analogues = da.get("analogues") or []
    if analogues:
        document.add_paragraph().add_run("Real systems that failed in this space (read before submission):").bold = True
        for a in analogues:
            item = document.add_paragraph(style="List Bullet")
            _set_run(item.add_run(f"{_text(a.get('title'))} ({_text(a.get('year'))}, {_text(a.get('source_name'))}): "), bold=True)
            item.add_run(f"{_text(a.get('lesson'))} {_text(a.get('source_url'))}")

    patterns = da.get("patterns") or []
    if patterns:
        document.add_paragraph().add_run("Patterns researchers like you often miss:").bold = True
        for p in patterns:
            item = document.add_paragraph(style="List Bullet")
            _set_run(item.add_run(f"{_text(p.get('pattern'))} "), bold=True)
            item.add_run(f"Usually missed: {_text(p.get('usually_missed'))} Ask: {_text(p.get('prompt'))}")

    bridge = da.get("real_evidence_bridge")
    if bridge:
        document.add_paragraph().add_run("Convert synthetic probes into real evidence:").bold = True
        document.add_paragraph(_text(bridge.get("notice")))
        for t in bridge.get("evidence_types") or []:
            item = document.add_paragraph(style="List Bullet")
            _set_run(item.add_run(f"{_text(t.get('label'))}: "), bold=True)
            item.add_run(_text(t.get("description")))


def build_committee_application_docx(session: Dict[str, Any]) -> bytes:
    """Return a DOCX ethics-application draft for ``session`` as raw bytes."""

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_paragraph()
    _set_run(
        title.add_run("ETHICS APPLICATION DRAFT — Generated by SafeBARS Ethical Mirror"),
        size=18, bold=True, color=INK,
    )
    subtitle = document.add_paragraph()
    _set_run(
        subtitle.add_run(_text(session.get("title"), "Untitled research project")),
        size=13, color=MUTED,
    )

    notice = document.add_paragraph()
    _set_run(notice.add_run(
        "DRAFT ONLY — this is not an approval or compliance verdict. Transfer "
        "the content below into your institution's required form. Generated or "
        "scaffolded text requires researcher verification and formal "
        "institutional ethics review."
    ), bold=True, color=AMBER)

    # 1. Research summary ---------------------------------------------------
    document.add_heading("1. Research summary", level=1)
    research_plan = _text(session.get("research_plan"))
    if research_plan:
        _label_paragraph(document, "Research plan (as submitted)", research_plan)
    intake = session.get("intake_answers") or {}
    for key in INTAKE_LABELS:
        if _text(intake.get(key)):
            _label_paragraph(document, INTAKE_LABELS[key], intake[key])
    commitments = session.get("value_commitments") or []
    if commitments:
        _label_paragraph(document, "Value commitments", "; ".join(commitments))

    # 2. Ethical considerations across literature lenses -------------------
    document.add_heading("2. Ethical considerations across literature lenses", level=1)
    lenses = session.get("lenses") or []
    if not lenses:
        document.add_paragraph(
            "Run the Mirror analysis to populate literature-grounded ethical "
            "dimensions. Each dimension below will show its coverage, the "
            "plan passage it was grounded in, and the relevant source."
        )
    for lens in lenses:
        heading = document.add_heading(level=2)
        _set_run(
            heading.add_run(_text(lens.get("label") or lens.get("id"), "Lens")),
            size=13, bold=True,
        )
        _label_paragraph(document, "Coverage", lens.get("status"))
        _label_paragraph(
            document,
            "Evidence / interpretation",
            _text(lens.get("evidence") or lens.get("interpretation")),
        )
        source_ids = lens.get("source_ids") or []
        if source_ids:
            _label_paragraph(
                document, "Literature source(s)",
                ", ".join(_text(s) for s in source_ids),
            )
        concern = lens.get("concern")
        if concern:
            _label_paragraph(
                document, "Why this matters for your project", _text(concern),
            )
        reflection_question = lens.get("reflection_question")
        if reflection_question:
            _label_paragraph(
                document, "Ask yourself", _text(reflection_question),
            )
        if lens.get("status") in WEAK_LENS_STATUSES:
            blank = document.add_paragraph()
            _set_run(blank.add_run("Researcher note (to complete): "), bold=True)
            blank.add_run("_" * 60)

    # 3. Identified risks and mitigations ----------------------------------
    document.add_heading("3. Identified risks and mitigations", level=1)
    scenarios = session.get("scenarios") or []
    if scenarios:
        document.add_paragraph(
            "The Mirror probed bounded roles and breakdown scenarios. Each is a "
            "hypothesis about how the plan could go wrong, not stakeholder "
            "testimony."
        )
        for scenario in scenarios:
            item = document.add_paragraph(style="List Bullet")
            label = _text(scenario.get("label") or scenario.get("title") or scenario.get("role"))
            desc = _text(scenario.get("description") or scenario.get("harm") or scenario.get("rationale"))
            _set_run(item.add_run(f"{label}: "), bold=True)
            item.add_run(desc)
    else:
        document.add_paragraph("No role probes were generated for this plan.")

    edges = session.get("dissonance_edges") or []
    if edges:
        document.add_paragraph("Tension paths the researcher must reconcile:")
        for edge in edges:
            item = document.add_paragraph(style="List Bullet")
            label = _text(edge.get("label") or edge.get("id"))
            status = _text(edge.get("status"))
            rationale = _text(edge.get("rationale") or edge.get("description"))
            _set_run(item.add_run(f"{label} [{status}]: "), bold=True)
            item.add_run(rationale)

    # 4. Deep-audit discovery cues -----------------------------------------
    _add_deep_audit_section(document, session)

    # 5. ESR societal & community risk statement ---------------------------
    _add_esr_section(document, session)

    # 6. Researcher revisions and decisions --------------------------------
    document.add_heading("6. Researcher revisions and decisions", level=1)
    revisions = session.get("revisions") or []
    if revisions:
        for revision in revisions:
            item = document.add_paragraph(style="List Bullet")
            summary = _text(
                revision.get("summary") or revision.get("revised_plan") or revision.get("note")
            )
            _set_run(item.add_run("Revision: "), bold=True)
            item.add_run(summary)
    else:
        document.add_paragraph(
            "No revisions recorded yet. Responding to tensions in Step 5 "
            "(Revise / Add safeguard / Contest / Consult) will populate this "
            "section and the change ledger."
        )

    # 7. References ---------------------------------------------------------
    document.add_heading("7. References", level=1)
    references = [
        "Bernstein, M. S., Levi, M., Magnus, D., Rajala, B. A., Satz, D., & "
        "Waeiss, C. (2021). Ethics and society review: Ethics reflection as a "
        "precondition to research funding. PNAS, 118(52), e2117261118. "
        "DOI 10.1073/pnas.2117261118.",
        "National Commission for the Protection of Human Subjects of "
        "Biomedical and Behavioral Research (1979). The Belmont Report: "
        "Ethical Principles and Guidelines for the Protection of Human "
        "Subjects of Research.",
        "Dittrich, D., Kenneally, E., et al. (2012). The Menlo Report: "
        "Ethical Principles Guiding Information and Communication Technology "
        "Research.",
        "NIST (2023). AI Risk Management Framework (AI RMF 1.0). NIST AI 100-1.",
        "Friedman, B., Kahn, P. H., Borning, A., & Huldtgren, A. (2013). "
        "Value Sensitive Design and information systems. In Early Engagement "
        "and New Technologies: Opening up the Laboratory.",
    ]
    for ref in references:
        item = document.add_paragraph(style="List Bullet")
        item.add_run(ref)

    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()
