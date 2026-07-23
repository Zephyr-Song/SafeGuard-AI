"""Human-readable Word and PDF reports for SafeBARS encounter sessions."""

from __future__ import annotations

from collections import Counter
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    LongTable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .ethics_application import build_application_readiness


INK = "17202A"
MUTED = "66727F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GREEN = "246B49"
TEAL = "176F73"
LIGHT_GRAY = "F2F4F7"
LINE = "D9DEE3"
AMBER_FILL = "FFF8E8"


def _text(value: Any, fallback: str = "Not provided") -> str:
    cleaned = "" if value is None else str(value).strip()
    return cleaned or fallback


def _date(value: str) -> str:
    if not value:
        return "Not recorded"
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
        return parsed.strftime("%Y-%m-%d %H:%M UTC")
    except ValueError:
        return value


def _passage_lookup(session: Dict[str, Any]) -> Dict[str, Dict[str, str]]:
    return {item["id"]: item for item in session.get("passages", [])}


def _decision_summary(session: Dict[str, Any]) -> str:
    counts = Counter(item.get("decision", "pending") for item in session.get("issues", []))
    order = ["accept", "edit", "reject", "defer", "pending"]
    return ", ".join(f"{name}: {counts[name]}" for name in order if counts[name]) or "No issues"


def _set_run(run, *, size: float = 11, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: List[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", indent_dxa)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _add_label_paragraph(doc: Document, label: str, value: Any, *, after: float = 4) -> Any:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    label_run = paragraph.add_run(f"{label}: ")
    _set_run(label_run, bold=True, color=DARK_BLUE)
    value_run = paragraph.add_run(_text(value))
    _set_run(value_run)
    return paragraph


def _configure_docx(
    document: Document,
    header_text: str = "SAFEBARS  /  ENCOUNTER STRESS-TEST REPORT",
) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = header_text
    _set_run(header.runs[0], size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    _add_page_number(footer)


def _add_docx_boundary_callout(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(8)
    paragraph.paragraph_format.right_indent = Pt(8)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.10
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), AMBER_FILL)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "D6B86A")
    borders.append(left)
    p_pr.append(borders)
    label = paragraph.add_run("Interpretation boundary. ")
    _set_run(label, bold=True, color="7A5A00")
    copy = paragraph.add_run(
        "This report contains planning hypotheses about the submitted protocol. It is not participant "
        "evidence, community representation, an ethics approval, or a prediction of real behavior."
    )
    _set_run(copy, color="5E4A17")


def build_docx_report(session: Dict[str, Any]) -> bytes:
    document = Document()
    _configure_docx(document)
    passages = _passage_lookup(session)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("SAFEBARS FULL AUDIT REPORT")
    _set_run(title_run, size=23, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(_text(session.get("project", {}).get("title"), "Untitled fieldwork plan"))
    _set_run(subtitle_run, size=14, color=MUTED)

    metadata_rows = [
        ("Session", session.get("id")),
        ("Generated", _date(session.get("updated_at", ""))),
        ("Status", session.get("status", "unknown")),
        (
            "Contents",
            f"{len(session.get('encounter_map', []))} stages | {len(session.get('traces', []))} traces | "
            f"{len(session.get('issues', []))} issues | {len(session.get('handoffs', []))} handoffs",
        ),
    ]
    for label, value in metadata_rows:
        metadata_line = document.add_paragraph()
        metadata_line.paragraph_format.space_after = Pt(4)
        metadata_line.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))
        _set_run(metadata_line.add_run(label), bold=True, color=DARK_BLUE)
        _set_run(metadata_line.add_run(f"\t{_text(value)}"))
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    _add_docx_boundary_callout(document)

    document.add_heading("1. Project context", level=1)
    _add_label_paragraph(
        document,
        "Research area and ethics-review context",
        session.get("project", {}).get("review_context"),
    )
    _add_label_paragraph(document, "Project plan", session.get("project", {}).get("context"))
    _add_label_paragraph(document, "People and relationships", session.get("project", {}).get("target_people"))
    _add_label_paragraph(document, "Decision summary", _decision_summary(session))

    document.add_heading("2. Encounter map", level=1)
    map_table = document.add_table(rows=1, cols=4)
    map_table.style = "Table Grid"
    _set_table_geometry(map_table, [2600, 1300, 1200, 4260])
    headers = ["Stage", "Coverage", "Scope", "Sources / responsibility note"]
    for cell, header_text in zip(map_table.rows[0].cells, headers):
        _set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        _set_run(paragraph.add_run(header_text), size=9.5, bold=True, color=DARK_BLUE)
    _set_repeat_table_header(map_table.rows[0])
    for stage in session.get("encounter_map", []):
        row = map_table.add_row()
        values = [
            stage.get("name"),
            stage.get("coverage"),
            "Included" if stage.get("included", True) else "Excluded",
            ", ".join(stage.get("source_passage_ids", []))
            + (f" | {stage.get('notes')}" if stage.get("notes") else ""),
        ]
        for cell, value in zip(row.cells, values):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            _set_run(paragraph.add_run(_text(value, "None")), size=9.5)
    _set_table_geometry(map_table, [2600, 1300, 1200, 4260])

    document.add_heading("3. Breakdown traces", level=1)
    if not session.get("traces"):
        document.add_paragraph("No scenario traces were run.")
    for index, trace in enumerate(session.get("traces", []), start=1):
        document.add_heading(f"3.{index} {trace.get('title', 'Untitled trace')}", level=2)
        _add_label_paragraph(document, "Status", trace.get("status"))
        for step in trace.get("steps", []):
            sources = ", ".join(step.get("source_passage_ids", [])) or "No cited passage"
            _add_label_paragraph(
                document,
                step.get("label", "Trace step"),
                f"{_text(step.get('text'))} [Sources: {sources}]",
            )
        _add_label_paragraph(document, "First gap or coverage result", trace.get("first_gap"))
        _add_label_paragraph(document, "Boundary", trace.get("uncertainty"), after=8)

    document.add_page_break()
    document.add_heading("4. Issue ledger", level=1)
    if not session.get("issues"):
        document.add_paragraph("No contestable issues were produced in this session.")
    for index, issue in enumerate(session.get("issues", []), start=1):
        document.add_heading(f"4.{index} {issue.get('title', 'Untitled issue')}", level=2)
        _add_label_paragraph(
            document,
            "Assessment",
            f"{_text(issue.get('severity')).upper()} | {_text(issue.get('decision')).upper()} | {_text(issue.get('agent'))}",
        )
        _add_label_paragraph(document, "Observation", issue.get("observation"))
        for passage_id in issue.get("source_passage_ids", []):
            passage = passages.get(passage_id)
            if passage:
                _add_label_paragraph(
                    document,
                    f"Source {passage_id} - {passage.get('artifact_label', '')}",
                    passage.get("text"),
                )
        _add_label_paragraph(document, "Proposed change", issue.get("suggestion"))
        _add_label_paragraph(document, "Researcher revision", issue.get("revised_text", ""))
        _add_label_paragraph(document, "Decision rationale", issue.get("decision_rationale", ""))
        _add_label_paragraph(document, "Boundary", issue.get("uncertainty"), after=8)
        for position in issue.get("agent_positions", []):
            _add_label_paragraph(
                document,
                f"Position - {position.get('agent', 'Agent')}",
                position.get("position"),
            )
        _add_label_paragraph(document, "Resolution rule", issue.get("resolution_rule"), after=8)

    document.add_heading("5. Real-world handoffs", level=1)
    if not session.get("handoffs"):
        document.add_paragraph("No real-world handoffs were recorded.")
    for index, handoff in enumerate(session.get("handoffs", []), start=1):
        document.add_heading(f"5.{index} {handoff.get('question', 'Unresolved question')}", level=2)
        _add_label_paragraph(
            document,
            "Priority / status",
            f"{_text(handoff.get('priority')).upper()} | {_text(handoff.get('status')).upper()}",
        )
        _add_label_paragraph(document, "Why AI cannot resolve this", handoff.get("why_ai_cannot_resolve"))
        _add_label_paragraph(document, "Recommended real-world owner", handoff.get("recommended_role_label", handoff.get("owner")))
        _add_label_paragraph(document, "Expert advice", handoff.get("expert_advice", ""))
        _add_label_paragraph(document, "Expert rationale", handoff.get("expert_rationale", ""), after=8)
        researcher_response = _add_label_paragraph(
            document,
            "Researcher response",
            handoff.get("researcher_response", ""),
        )
        researcher_response.paragraph_format.keep_with_next = True
        _add_label_paragraph(document, "Linked protocol revision", handoff.get("researcher_revised_text", ""), after=8)

    document.add_heading("Appendix A. Framework-grounded ethics map", level=1)
    assessment = session.get("framework_assessment", {})
    _add_label_paragraph(document, "Pathway", assessment.get("pathway"))
    _add_label_paragraph(
        document,
        "Active frameworks",
        ", ".join(item.get("name", "") for item in assessment.get("frameworks", [])),
    )
    for dimension in assessment.get("dimensions", []):
        document.add_heading(f"{dimension.get('label', 'Framework dimension')} - {_text(dimension.get('coverage')).upper()}", level=2)
        _add_label_paragraph(document, "Framework question", dimension.get("question"))
        _add_label_paragraph(document, "Submitted evidence", ", ".join(dimension.get("source_passage_ids", [])), after=8)

    document.add_heading("Appendix B. Inspectable audit plan", level=1)
    for index, task in enumerate(session.get("audit_plan", []), start=1):
        document.add_heading(f"B.{index} {task.get('title', 'Untitled task')}", level=2)
        _add_label_paragraph(
            document,
            "Agent / state",
            f"{_text(task.get('agent'))} | {_text(task.get('priority')).upper()} | {_text(task.get('status')).upper()} | attempt {task.get('attempts', 0)}",
        )
        _add_label_paragraph(document, "Routing reason", task.get("reason"))
        _add_label_paragraph(document, "Inputs", ", ".join(task.get("input_passage_ids", [])))
        _add_label_paragraph(document, "Tools", ", ".join(task.get("tools", [])))
        _add_label_paragraph(document, "Dependencies", ", ".join(task.get("depends_on", [])))
        _add_label_paragraph(document, "Stop condition", task.get("stop_condition"))
        _add_label_paragraph(document, "Result", task.get("result_summary"), after=8)

    document.add_heading("Appendix C. Submitted materials", level=1)
    artifact_labels = {
        "recruitment": "Recruitment message",
        "consent": "Consent language",
        "interview": "Interview questions",
        "activity": "Workshop or activity plan",
        "safety": "Safety and escalation procedure",
        "follow_up": "Debrief, follow-up, and data use",
    }
    for key, label in artifact_labels.items():
        document.add_heading(label, level=2)
        document.add_paragraph(_text(session.get("artifacts", {}).get(key)))

    document.add_heading("Appendix D. Audit event history", level=1)
    for event in session.get("event_log", []):
        _add_label_paragraph(
            document,
            _date(event.get("created_at", "")),
            f"{event.get('event_type', 'event')} - {event.get('payload', {})}",
        )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_styles() -> Dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "SafeBARSTitle", parent=base["Title"], fontName="Helvetica-Bold", fontSize=23,
            leading=27, textColor=colors.HexColor(f"#{INK}"), alignment=TA_LEFT, spaceAfter=4,
        ),
        "subtitle": ParagraphStyle(
            "SafeBARSSubtitle", parent=base["Normal"], fontName="Helvetica", fontSize=14,
            leading=18, textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=14,
        ),
        "body": ParagraphStyle(
            "SafeBARSBody", parent=base["BodyText"], fontName="Helvetica", fontSize=9.5,
            leading=12.5, textColor=colors.HexColor(f"#{INK}"), spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "SafeBARSSmall", parent=base["BodyText"], fontName="Helvetica", fontSize=8.2,
            leading=10.5, textColor=colors.HexColor(f"#{INK}"), spaceAfter=2,
        ),
        "h1": ParagraphStyle(
            "SafeBARSH1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=16,
            leading=19, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=16, spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "SafeBARSH2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12,
            leading=15, textColor=colors.HexColor(f"#{DARK_BLUE}"), spaceBefore=10, spaceAfter=5,
        ),
        "boundary": ParagraphStyle(
            "SafeBARSBoundary", parent=base["BodyText"], fontName="Helvetica", fontSize=9.5,
            leading=13, textColor=colors.HexColor("#5E4A17"), spaceAfter=0,
        ),
    }


def _p(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(_text(value)).replace("\n", "<br/>"), style)


def _label_pdf(label: str, value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(f"<b>{escape(label)}:</b> {escape(_text(value)).replace(chr(10), '<br/>')}", style)


def _pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor(f"#{LINE}"))
    canvas.setLineWidth(0.5)
    canvas.line(doc.leftMargin, letter[1] - 0.58 * inch, letter[0] - doc.rightMargin, letter[1] - 0.58 * inch)
    canvas.setFont("Helvetica-Bold", 8)
    canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
    canvas.drawString(doc.leftMargin, letter[1] - 0.48 * inch, "SAFEBARS / ENCOUNTER STRESS-TEST REPORT")
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(letter[0] - doc.rightMargin, 0.5 * inch, f"Page {doc.page}")
    canvas.restoreState()


def _pdf_meta_table(session: Dict[str, Any], styles: Dict[str, ParagraphStyle]) -> Table:
    data = [
        [_p("Session", styles["small"]), _p(session.get("id"), styles["small"])],
        [_p("Generated", styles["small"]), _p(_date(session.get("updated_at", "")), styles["small"])],
        [_p("Status", styles["small"]), _p(session.get("status"), styles["small"])],
        [
            _p("Contents", styles["small"]),
            _p(
                f"{len(session.get('encounter_map', []))} stages | {len(session.get('traces', []))} traces | "
                f"{len(session.get('issues', []))} issues | {len(session.get('handoffs', []))} handoffs",
                styles["small"],
            ),
        ],
    ]
    table = Table(data, colWidths=[1.25 * inch, 5.25 * inch], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{LIGHT_GRAY}")),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor(f"#{DARK_BLUE}")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{LINE}")),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def build_pdf_report(session: Dict[str, Any]) -> bytes:
    output = BytesIO()
    styles = _pdf_styles()
    passages = _passage_lookup(session)
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.82 * inch,
        bottomMargin=0.72 * inch,
        title="SafeBARS Encounter Stress-Test Report",
        author="SafeBARS",
    )
    story: List[Any] = []
    story.append(Spacer(1, 0.12 * inch))
    story.append(Paragraph("SAFEBARS FULL AUDIT REPORT", styles["title"]))
    story.append(_p(session.get("project", {}).get("title", "Untitled fieldwork plan"), styles["subtitle"]))
    story.append(_pdf_meta_table(session, styles))
    story.append(Spacer(1, 0.14 * inch))
    boundary = Table(
        [[Paragraph(
            "<b>Interpretation boundary.</b> This report contains planning hypotheses about the submitted "
            "protocol. It is not participant evidence, community representation, an ethics approval, or a "
            "prediction of real behavior.",
            styles["boundary"],
        )]],
        colWidths=[6.5 * inch],
    )
    boundary.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{AMBER_FILL}")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#D6B86A")),
        ("LEFTPADDING", (0, 0), (-1, -1), 9),
        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(boundary)

    story.append(Paragraph("1. Project context", styles["h1"]))
    story.append(
        _label_pdf(
            "Research area and ethics-review context",
            session.get("project", {}).get("review_context"),
            styles["body"],
        )
    )
    story.append(_label_pdf("Project plan", session.get("project", {}).get("context"), styles["body"]))
    story.append(_label_pdf("People and relationships", session.get("project", {}).get("target_people"), styles["body"]))
    story.append(_label_pdf("Decision summary", _decision_summary(session), styles["body"]))

    story.append(Paragraph("2. Encounter map", styles["h1"]))
    map_rows = [[
        _p("Stage", styles["small"]), _p("Coverage", styles["small"]),
        _p("Scope", styles["small"]), _p("Sources / responsibility note", styles["small"]),
    ]]
    for stage in session.get("encounter_map", []):
        sources = ", ".join(stage.get("source_passage_ids", []))
        note = f"{sources} | {stage.get('notes')}" if stage.get("notes") else sources
        map_rows.append([
            _p(stage.get("name"), styles["small"]),
            _p(stage.get("coverage"), styles["small"]),
            _p("Included" if stage.get("included", True) else "Excluded", styles["small"]),
            _p(note or "None", styles["small"]),
        ])
    map_table = LongTable(map_rows, colWidths=[1.8 * inch, 0.9 * inch, 0.8 * inch, 3.0 * inch], repeatRows=1)
    map_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_GRAY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{DARK_BLUE}")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{LINE}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(map_table)

    story.append(Paragraph("3. Breakdown traces", styles["h1"]))
    if not session.get("traces"):
        story.append(_p("No scenario traces were run.", styles["body"]))
    for index, trace in enumerate(session.get("traces", []), start=1):
        block: List[Any] = [Paragraph(f"3.{index} {escape(_text(trace.get('title')))}", styles["h2"])]
        block.append(_label_pdf("Status", trace.get("status"), styles["body"]))
        for step in trace.get("steps", []):
            sources = ", ".join(step.get("source_passage_ids", [])) or "No cited passage"
            block.append(_label_pdf(step.get("label", "Trace step"), f"{_text(step.get('text'))} [Sources: {sources}]", styles["body"]))
        block.append(_label_pdf("First gap or coverage result", trace.get("first_gap"), styles["body"]))
        block.append(_label_pdf("Boundary", trace.get("uncertainty"), styles["body"]))
        story.extend(block)

    story.append(PageBreak())
    story.append(Paragraph("4. Issue ledger", styles["h1"]))
    if not session.get("issues"):
        story.append(_p("No contestable issues were produced in this session.", styles["body"]))
    for index, issue in enumerate(session.get("issues", []), start=1):
        story.append(Paragraph(f"4.{index} {escape(_text(issue.get('title')))}", styles["h2"]))
        story.append(_label_pdf(
            "Assessment",
            f"{_text(issue.get('severity')).upper()} | {_text(issue.get('decision')).upper()} | {_text(issue.get('agent'))}",
            styles["body"],
        ))
        story.append(_label_pdf("Observation", issue.get("observation"), styles["body"]))
        for passage_id in issue.get("source_passage_ids", []):
            passage = passages.get(passage_id)
            if passage:
                source_box = Table(
                    [[_label_pdf(f"Source {passage_id} - {passage.get('artifact_label', '')}", passage.get("text"), styles["small"])]],
                    colWidths=[6.5 * inch],
                )
                source_box.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{LIGHT_GRAY}")),
                    ("LINEBEFORE", (0, 0), (0, -1), 1.2, colors.HexColor(f"#{BLUE}")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(source_box)
                story.append(Spacer(1, 3))
        story.append(_label_pdf("Proposed change", issue.get("suggestion"), styles["body"]))
        story.append(_label_pdf("Researcher revision", issue.get("revised_text", ""), styles["body"]))
        story.append(_label_pdf("Decision rationale", issue.get("decision_rationale", ""), styles["body"]))
        story.append(_label_pdf("Boundary", issue.get("uncertainty"), styles["body"]))
        for position in issue.get("agent_positions", []):
            story.append(_label_pdf(
                f"Position - {position.get('agent', 'Agent')}",
                position.get("position"),
                styles["body"],
            ))
        story.append(_label_pdf("Resolution rule", issue.get("resolution_rule"), styles["body"]))

    story.append(Paragraph("5. Real-world handoffs", styles["h1"]))
    if not session.get("handoffs"):
        story.append(_p("No real-world handoffs were recorded.", styles["body"]))
    for index, handoff in enumerate(session.get("handoffs", []), start=1):
        story.append(Paragraph(f"5.{index} {escape(_text(handoff.get('question')))}", styles["h2"]))
        story.append(_label_pdf(
            "Priority / status",
            f"{_text(handoff.get('priority')).upper()} | {_text(handoff.get('status')).upper()}",
            styles["body"],
        ))
        story.append(_label_pdf("Why AI cannot resolve this", handoff.get("why_ai_cannot_resolve"), styles["body"]))
        story.append(_label_pdf("Recommended real-world owner", handoff.get("recommended_role_label", handoff.get("owner")), styles["body"]))
        story.append(_label_pdf("Expert advice", handoff.get("expert_advice", ""), styles["body"]))
        story.append(_label_pdf("Expert rationale", handoff.get("expert_rationale", ""), styles["body"]))

    story.append(PageBreak())
    story.append(Paragraph("Appendix A. Framework-grounded ethics map", styles["h1"]))
    assessment = session.get("framework_assessment", {})
    story.append(_label_pdf("Pathway", assessment.get("pathway"), styles["body"]))
    story.append(_label_pdf(
        "Active frameworks",
        ", ".join(item.get("name", "") for item in assessment.get("frameworks", [])),
        styles["body"],
    ))
    for dimension in assessment.get("dimensions", []):
        story.append(Paragraph(
            f"{escape(_text(dimension.get('label')))} - {escape(_text(dimension.get('coverage')).upper())}",
            styles["h2"],
        ))
        story.append(_label_pdf("Framework question", dimension.get("question"), styles["body"]))
        story.append(_label_pdf("Submitted evidence", ", ".join(dimension.get("source_passage_ids", [])), styles["body"]))

    story.append(Paragraph("Appendix B. Inspectable audit plan", styles["h1"]))
    for index, task in enumerate(session.get("audit_plan", []), start=1):
        story.append(Paragraph(f"B.{index} {escape(_text(task.get('title')))}", styles["h2"]))
        story.append(_label_pdf(
            "Agent / state",
            f"{_text(task.get('agent'))} | {_text(task.get('priority')).upper()} | {_text(task.get('status')).upper()} | attempt {task.get('attempts', 0)}",
            styles["body"],
        ))
        story.append(_label_pdf("Routing reason", task.get("reason"), styles["body"]))
        story.append(_label_pdf("Inputs", ", ".join(task.get("input_passage_ids", [])), styles["body"]))
        story.append(_label_pdf("Tools", ", ".join(task.get("tools", [])), styles["body"]))
        story.append(_label_pdf("Dependencies", ", ".join(task.get("depends_on", [])), styles["body"]))
        story.append(_label_pdf("Stop condition", task.get("stop_condition"), styles["body"]))
        story.append(_label_pdf("Result", task.get("result_summary"), styles["body"]))

    story.append(Paragraph("Appendix C. Submitted materials", styles["h1"]))
    artifact_labels = {
        "recruitment": "Recruitment message",
        "consent": "Consent language",
        "interview": "Interview questions",
        "activity": "Workshop or activity plan",
        "safety": "Safety and escalation procedure",
        "follow_up": "Debrief, follow-up, and data use",
    }
    for key, label in artifact_labels.items():
        story.append(Paragraph(label, styles["h2"]))
        story.append(_p(session.get("artifacts", {}).get(key), styles["body"]))

    story.append(Paragraph("Appendix D. Audit event history", styles["h1"]))
    for event in session.get("event_log", []):
        story.append(_label_pdf(
            _date(event.get("created_at", "")),
            f"{event.get('event_type', 'event')} - {event.get('payload', {})}",
            styles["small"],
        ))

    document.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    return output.getvalue()


def _add_esr_societal_risk_section(document, assessment):
    """Add an ESR-style societal/community risk statement to the committee draft.

    Operationalizes Bernstein et al. (2021), PNAS: beyond the individual-participant
    focus of a standard IRB, researchers should state risks to society, to subgroups,
    and globally, and commit to mitigation strategies. SafeBARS surfaces the societal
    dimensions it already assessed and leaves the mitigation statements for the
    researcher to verify and complete.
    """
    document.add_paragraph(
        "Standard ethics review (for example an IRB under the Common Rule) focuses on risks to "
        "individual participants. The Ethics and Society Review (ESR; Bernstein, Levi, Magnus, "
        "Rajala, Satz & Waeiss, PNAS 2021, DOI 10.1073/pnas.2117261118) argues that research can "
        "also create risks to society, to subgroups within society, and globally, and that "
        "researchers should state those risks and commit to mitigation strategies. SafeBARS "
        "surfaced the considerations below from your framework assessment. Verify and complete "
        "the mitigation statements, then map them to the institution's current form in "
        "consultation with the responsible committee. This statement supplements, and does not "
        "replace, individual-participant review."
    )
    societal_frameworks = {
        "menlo",
        "ai_irb_questions",
        "ai_rec_guidance",
        "nist_ai_rmf",
        "vsd",
        "esr",
    }
    dimensions = [
        d for d in assessment.get("dimensions", [])
        if d.get("framework") in societal_frameworks
    ]
    if not dimensions:
        document.add_paragraph(
            "No societal, ICT, AI-governance, or stakeholder-value dimensions were activated for "
            "this protocol. Even so, state any foreseeable risks to subgroups or to society and "
            "your planned mitigations in the space below."
        )
    for dim in dimensions:
        document.add_heading(
            f"{dim.get('label', 'Consideration')}  [{dim.get('framework')}, {_text(dim.get('coverage')).upper()}]",
            level=2,
        )
        document.add_paragraph(_text(dim.get("question")))
        document.add_paragraph(
            f"Coverage evidence: {', '.join(dim.get('source_passage_ids', [])) or 'none located'}."
        )
        blank = document.add_paragraph()
        _set_run(blank.add_run("Mitigation statement (researcher to complete): "), bold=True)
        blank.add_run("_" * 60)


def build_ethics_application_docx(session: Dict[str, Any]) -> bytes:
    """Build a generic application draft, never an approval or compliance verdict."""
    document = Document()
    _configure_docx(document, "SAFEBARS  /  ETHICS APPLICATION DRAFT")
    project = session.get("project", {})
    artifacts = session.get("artifacts", {})
    assessment = session.get("framework_assessment", {})
    readiness = session.get("application_readiness") or build_application_readiness(session)

    title = document.add_paragraph()
    _set_run(title.add_run("ETHICS APPLICATION DRAFT"), size=22, color=INK, bold=True)
    subtitle = document.add_paragraph()
    _set_run(subtitle.add_run(_text(project.get("title"), "Untitled research project")), size=14, color=MUTED)
    _add_docx_boundary_callout(document)
    warning = document.add_paragraph()
    _set_run(
        warning.add_run(
            "This generic draft must be transferred into the current form required by the relevant institution. "
            "Generated or scaffolded text requires researcher verification and formal institutional review."
        ),
        color="7A5A00",
        bold=True,
    )

    document.add_heading("Draft profile and completeness", level=1)
    _add_label_paragraph(document, "Application profile", readiness.get("profile", {}).get("label"))
    _add_label_paragraph(document, "Documented fields", f"{readiness.get('counts', {}).get('documented', 0)} of {len(readiness.get('fields', []))}")
    _add_label_paragraph(document, "Unresolved handoffs", readiness.get("unresolved_handoff_count", 0))
    for field in readiness.get("fields", []):
        document.add_heading(
            f"{field.get('label', 'Application field')} - {_text(field.get('status')).upper()}",
            level=2,
        )
        _add_label_paragraph(
            document,
            "Evidence sources",
            ", ".join(item.get("source", "") for item in field.get("evidence", [])),
        )
        if field.get("status") != "documented":
            _add_label_paragraph(document, "Refinement needed", field.get("prompt"), after=8)

    document.add_heading("1. Application overview", level=1)
    _add_label_paragraph(document, "Project title", project.get("title"))
    _add_label_paragraph(
        document,
        "Research area and ethics-review context",
        project.get("review_context"),
    )
    _add_label_paragraph(document, "Research context and purpose", project.get("context"))
    _add_label_paragraph(document, "Participants and affected relationships", project.get("target_people"))
    _add_label_paragraph(document, "Framework pathway", assessment.get("pathway"))
    _add_label_paragraph(document, "Research builds, studies, or uses AI", "Yes" if project.get("uses_ai") else "No")

    document.add_heading("2. Recruitment and participant selection", level=1)
    document.add_paragraph(_text(artifacts.get("recruitment")))
    _add_label_paragraph(document, "Belmont justice check", "Justify necessary eligibility boundaries, fair access, burdens, benefits, gatekeepers, and compensation.")

    document.add_heading("3. Participant information, consent, and withdrawal", level=1)
    document.add_paragraph(_text(artifacts.get("consent")))
    _add_label_paragraph(document, "Belmont respect check", "Verify information, comprehension, voluntariness, recording choices, question skipping, and withdrawal at each data stage.")

    document.add_heading("4. Study procedures and participant activities", level=1)
    _add_label_paragraph(document, "Interview questions or prompts", artifacts.get("interview"))
    _add_label_paragraph(document, "Workshop or activity plan", artifacts.get("activity"))

    document.add_heading("5. Foreseeable risk, safeguarding, and support", level=1)
    _add_label_paragraph(document, "Safety and escalation procedure", artifacts.get("safety"))
    _add_label_paragraph(document, "Debrief and follow-up", artifacts.get("follow_up"))
    _add_label_paragraph(document, "Belmont beneficence check", "Explain how harms are minimized, benefits justified, and responsible staff, escalation limits, and support routes confirmed.")

    document.add_heading("6. Data management and confidentiality", level=1)
    _add_label_paragraph(document, "Submitted data-use plan", artifacts.get("follow_up"))
    _add_label_paragraph(document, "Required verification", "Specify data categories, collection, access, transfer, storage, security, retention, deletion, withdrawal effects, quotations, and reporting.")

    if project.get("uses_ai") or assessment.get("uses_ai"):
        document.add_heading("7. AI use and risk-management appendix", level=1)
        _add_label_paragraph(document, "AI role described by researcher", project.get("context"))
        _add_label_paragraph(
            document,
            "Submitted AI ethics-review supplement",
            artifacts.get("ai_governance"),
        )
        _add_label_paragraph(
            document,
            "Review basis",
            (
                "Makridis et al. (2023), AI-specific questions for human-subjects review "
                "(doi:10.3389/fcomp.2023.1235226); Connelly et al. (2025), guidance for "
                "research ethics committees and researchers in the age of AI "
                "(doi:10.5281/zenodo.13739834); and NIST AI RMF 1.0."
            ),
        )
        frameworks = {
            item.get("id"): item for item in assessment.get("frameworks", [])
        }
        for function in (
            "ai_govern",
            "ai_map",
            "ai_review_pathway",
            "ai_measure",
            "ai_manage",
        ):
            dimension = next((item for item in assessment.get("dimensions", []) if item.get("id") == function), None)
            if dimension:
                framework_name = frameworks.get(dimension.get("framework"), {}).get(
                    "name", dimension.get("framework")
                )
                _add_label_paragraph(
                    document,
                    f"{framework_name} - {dimension.get('label')}",
                    f"{_text(dimension.get('coverage')).upper()}: {dimension.get('question')} Evidence: {', '.join(dimension.get('source_passage_ids', [])) or 'none located'}",
                )
        next_section = 8
    else:
        next_section = 7

    # ESR societal and community risk statement beyond individual-subject review.
    document.add_heading(f"{next_section}. Societal & Community Risk Statement (Ethics and Society Review)", level=1)
    _add_esr_societal_risk_section(document, assessment)

    document.add_heading(f"{next_section + 1}. Researcher decisions and revisions", level=1)
    if not session.get("issues"):
        document.add_paragraph("No encounter-audit issues have been generated yet.")
    for issue in session.get("issues", []):
        document.add_heading(issue.get("title", "Protocol issue"), level=2)
        _add_label_paragraph(document, "Decision", issue.get("decision"))
        _add_label_paragraph(document, "Researcher rationale", issue.get("decision_rationale"))
        _add_label_paragraph(document, "Revised text", issue.get("revised_text"))

    document.add_heading(f"{next_section + 2}. Outstanding expert and stakeholder review", level=1)
    if not session.get("handoffs"):
        document.add_paragraph("No handoffs have been generated yet.")
    for handoff in session.get("handoffs", []):
        document.add_heading(handoff.get("question", "Unresolved question"), level=2)
        _add_label_paragraph(document, "Priority / status", f"{handoff.get('priority', 'standard')} / {handoff.get('status', 'open')}")
        _add_label_paragraph(document, "Recommended reviewer", handoff.get("recommended_role_label", handoff.get("owner")))
        _add_label_paragraph(document, "Why AI stopped", handoff.get("why_ai_cannot_resolve"))
        _add_label_paragraph(document, "Expert advice", handoff.get("expert_advice"))
        _add_label_paragraph(document, "Researcher response", handoff.get("researcher_response"))
        _add_label_paragraph(document, "Linked protocol revision", handoff.get("researcher_revised_text"))

    document.add_heading(f"{next_section + 3}. Submission readiness statement", level=1)
    document.add_paragraph(
        f"This draft is {readiness.get('completion_percent', 0)}% documented under the selected generic profile and "
        f"contains {readiness.get('unresolved_handoff_count', 0)} unresolved handoff(s). The researcher must verify every section, "
        "complete the institution-specific form, disclose AI assistance where required, and obtain formal approval before recruitment or data collection."
    )

    document.add_heading("References", level=1)
    document.add_paragraph(
        "Bernstein, M. S., Levi, M., Magnus, D., Rajala, B. A., Satz, D., & Waeiss, C. (2021). "
        "Ethics and society review: Ethics reflection as a precondition to research funding. "
        "Proceedings of the National Academy of Sciences, 118(52), e2117261118. "
        "https://doi.org/10.1073/pnas.2117261118"
    )
    document.add_paragraph(
        "Makridis, C. A., et al. (2023). Informing the ethical review of human subjects "
        "research utilizing artificial intelligence. Frontiers in Computer Science, 5, "
        "1235226. https://doi.org/10.3389/fcomp.2023.1235226"
    )
    document.add_paragraph(
        "Connelly, R., Osborne, N., Black, S., & Terras, M. (2025). Guidance for research "
        "ethics committees and researchers on designing research in the age of AI. "
        "https://doi.org/10.5281/zenodo.13739834"
    )
    document.add_paragraph(
        "U.S. National Commission for the Protection of Human Subjects (1979). The Belmont Report: "
        "Ethical Principles and Guidelines for the Protection of Human Subjects of Research."
    )
    document.add_paragraph(
        "Menlo Report (2012). Identifying Empirical Research Opportunities in Cybersecurity and "
        "Privacy; NIST AI Risk Management Framework (NIST AI 100-1); Value Sensitive Design "
        "(Friedman, Kahn & Borning)."
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_research_design_docx(session: Dict[str, Any]) -> bytes:
    """Build a researcher-facing study design grounded only in submitted materials."""
    document = Document()
    _configure_docx(document, "SAFEBARS  /  RESEARCH DESIGN")
    project = session.get("project", {})
    artifacts = session.get("artifacts", {})
    assessment = session.get("framework_assessment", {})
    readiness = session.get("application_readiness") or build_application_readiness(session)

    title = document.add_paragraph()
    _set_run(title.add_run("RESEARCH DESIGN AND ETHICS-IN-PRACTICE PLAN"), size=22, color=INK, bold=True)
    subtitle = document.add_paragraph()
    _set_run(subtitle.add_run(_text(project.get("title"), "Untitled research project")), size=14, color=MUTED)
    _add_docx_boundary_callout(document)
    document.add_paragraph(
        "This design plan organizes the researcher's submitted materials into a fieldwork-ready structure. "
        "It does not invent missing methods, replace disciplinary review, or constitute ethics approval."
    )

    document.add_heading("1. Study purpose, setting, and methodological rationale", level=1)
    _add_label_paragraph(
        document,
        "Research area and ethics-review context",
        project.get("review_context"),
    )
    _add_label_paragraph(document, "Submitted research context", project.get("context"))
    _add_label_paragraph(document, "Research aims or questions to verify", project.get("context"))
    _add_label_paragraph(document, "Framework pathway", assessment.get("pathway"))

    document.add_heading("2. Participants, relationships, sampling, and recruitment", level=1)
    _add_label_paragraph(document, "Participants and affected relationships", project.get("target_people"))
    _add_label_paragraph(document, "Recruitment and selection plan", artifacts.get("recruitment"))
    _add_label_paragraph(
        document,
        "Design check",
        "Verify inclusion and exclusion criteria, sampling rationale, access routes, gatekeepers, compensation, burden, and pressure safeguards.",
    )

    document.add_heading("3. Research setting, procedures, and participant journey", level=1)
    _add_label_paragraph(document, "Interview questions or prompts", artifacts.get("interview"))
    _add_label_paragraph(document, "Workshop, task, or activity plan", artifacts.get("activity"))
    stage_summary = ", ".join(
        f"{stage.get('name', 'Unnamed stage')} ({stage.get('coverage', 'unknown')})"
        for stage in session.get("encounter_map", [])
        if stage.get("included", True)
    )
    _add_label_paragraph(document, "Encounter stages to operationalize", stage_summary)
    _add_label_paragraph(
        document,
        "Implementation details to verify",
        "Confirm setting, duration, sequence, facilitator roles, accessibility, breaks, alternatives, recording, and dependencies between activities.",
    )

    document.add_heading("4. Consent, participation choices, and withdrawal", level=1)
    _add_label_paragraph(document, "Submitted consent procedure", artifacts.get("consent"))
    _add_label_paragraph(
        document,
        "Operational design check",
        "Specify when consent is revisited, how comprehension is checked, how questions may be skipped, and what withdrawal means before and after data collection.",
    )

    document.add_heading("5. Foreseeable risk, safeguarding, and support", level=1)
    _add_label_paragraph(document, "Safety and escalation procedure", artifacts.get("safety"))
    _add_label_paragraph(document, "Debrief, complaints, and follow-up", artifacts.get("follow_up"))
    _add_label_paragraph(
        document,
        "Roles and boundaries to verify",
        "Name responsible staff, pause and stop triggers, confidentiality limits, escalation thresholds, real support routes, and follow-up ownership.",
    )

    document.add_heading("6. Data generation, analysis, confidentiality, and reporting", level=1)
    _add_label_paragraph(document, "Submitted data and follow-up plan", artifacts.get("follow_up"))
    _add_label_paragraph(
        document,
        "Data lifecycle to verify",
        "Document data categories, capture, transcription, analysis approach, access, transfer, storage, security, retention, deletion, quotation, reporting, and withdrawal effects.",
    )

    if project.get("uses_ai") or assessment.get("uses_ai"):
        document.add_heading("7. AI role, human oversight, and failure response", level=1)
        _add_label_paragraph(document, "Submitted AI role", project.get("context"))
        _add_label_paragraph(
            document,
            "Submitted AI ethics-review supplement",
            artifacts.get("ai_governance"),
        )
        _add_label_paragraph(
            document,
            "Design basis",
            (
                "AI-specific human-subjects review questions (Makridis et al., 2023), "
                "university REC guidance for research in the age of AI (Connelly et al., "
                "2025), and NIST AI RMF 1.0."
            ),
        )
        frameworks = {
            item.get("id"): item for item in assessment.get("frameworks", [])
        }
        for function in (
            "ai_govern",
            "ai_map",
            "ai_review_pathway",
            "ai_measure",
            "ai_manage",
        ):
            dimension = next((item for item in assessment.get("dimensions", []) if item.get("id") == function), None)
            if dimension:
                framework_name = frameworks.get(dimension.get("framework"), {}).get(
                    "name", dimension.get("framework")
                )
                _add_label_paragraph(
                    document,
                    f"{framework_name} - {dimension.get('label')}",
                    f"{_text(dimension.get('coverage')).upper()}: {dimension.get('question')} Evidence: {', '.join(dimension.get('source_passage_ids', [])) or 'none located'}",
                )
        next_section = 8
    else:
        next_section = 7

    document.add_heading(f"{next_section}. Ethics-informed trade-offs and design decisions", level=1)
    tradeoffs = assessment.get("tradeoffs", [])
    if not tradeoffs:
        document.add_paragraph("No framework-grounded trade-off records are available yet.")
    deliberations = session.get("tradeoff_deliberations", {})
    for tradeoff in tradeoffs:
        document.add_heading(tradeoff.get("label", tradeoff.get("title", "Design trade-off")), level=2)
        left = tradeoff.get("left", {})
        right = tradeoff.get("right", {})
        deliberation = deliberations.get(tradeoff.get("id"), {})
        left_value = deliberation.get("value", left.get("value", 50))
        right_value = 100 - int(left_value)
        _add_label_paragraph(
            document,
            "Connected parameters",
            f"{_text(left.get('label'))} versus {_text(right.get('label'))}; framework dimensions: {', '.join(tradeoff.get('dimensions', [])) or 'not specified'}",
        )
        _add_label_paragraph(
            document,
            "Current design balance",
            f"{_text(left.get('label'))}: {left_value} / {_text(right.get('label'))}: {right_value}",
        )
        _add_label_paragraph(document, "Researcher decision prompt", tradeoff.get("prompt"))
        _add_label_paragraph(document, "Researcher rationale", deliberation.get("rationale"), after=8)

    document.add_heading(f"{next_section + 1}. Expert dependencies and unresolved design questions", level=1)
    unresolved = [item for item in session.get("handoffs", []) if item.get("status") != "resolved"]
    if not unresolved:
        document.add_paragraph("No unresolved expert or stakeholder handoffs are recorded.")
    for handoff in unresolved:
        document.add_heading(handoff.get("question", "Unresolved design question"), level=2)
        _add_label_paragraph(document, "Recommended reviewer", handoff.get("recommended_role_label", handoff.get("owner")))
        _add_label_paragraph(document, "Why AI stopped", handoff.get("why_ai_cannot_resolve"))
        _add_label_paragraph(document, "Expert advice", handoff.get("expert_advice"))
        _add_label_paragraph(document, "Researcher response", handoff.get("researcher_response"))
        _add_label_paragraph(document, "Linked protocol revision", handoff.get("researcher_revised_text"), after=8)

    document.add_heading(f"{next_section + 2}. Design readiness and next actions", level=1)
    _add_label_paragraph(document, "Application profile", readiness.get("profile", {}).get("label"))
    _add_label_paragraph(document, "Documented fields", f"{readiness.get('counts', {}).get('documented', 0)} of {len(readiness.get('fields', []))}")
    _add_label_paragraph(document, "Unresolved handoffs", readiness.get("unresolved_handoff_count", 0))
    for field in readiness.get("fields", []):
        if field.get("status") != "documented":
            _add_label_paragraph(document, field.get("label", "Incomplete design field"), field.get("prompt"))
    document.add_paragraph(
        "Before fieldwork, the researcher must confirm the final design with relevant methods, community, domain, data-governance, and institutional ethics reviewers."
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_expert_summary_docx(session: Dict[str, Any]) -> bytes:
    """Build a concise expert-facing queue and advice record for one protocol."""
    document = Document()
    _configure_docx(document, "SAFEBARS  /  EXPERT REVIEW SUMMARY")
    project = session.get("project", {})
    issue_by_id = {item["id"]: item for item in session.get("issues", [])}
    passages = _passage_lookup(session)

    title = document.add_paragraph()
    _set_run(title.add_run("SAFEBARS EXPERT REVIEW SUMMARY"), size=22, color=INK, bold=True)
    subtitle = document.add_paragraph()
    _set_run(subtitle.add_run(_text(project.get("title"), "Untitled research project")), size=14, color=MUTED)
    _add_docx_boundary_callout(document)
    _add_label_paragraph(document, "Project context", project.get("context"))
    _add_label_paragraph(document, "Framework pathway", session.get("framework_assessment", {}).get("pathway"))

    handoffs = sorted(
        session.get("handoffs", []),
        key=lambda item: ({"high": 0, "medium": 1, "standard": 2}.get(item.get("priority", "standard"), 3), item.get("status") == "resolved"),
    )
    document.add_heading("Prioritized handoff queue", level=1)
    if not handoffs:
        document.add_paragraph("No expert handoffs were generated for this session.")
    for index, handoff in enumerate(handoffs, start=1):
        issue = issue_by_id.get(handoff.get("issue_id"), {})
        document.add_heading(f"{index}. {issue.get('title', handoff.get('question', 'Unresolved issue'))}", level=2)
        _add_label_paragraph(document, "Priority / status", f"{handoff.get('priority', 'standard')} / {handoff.get('status', 'open')}")
        _add_label_paragraph(document, "Recommended reviewer", handoff.get("recommended_role_label", handoff.get("owner")))
        _add_label_paragraph(document, "Question", handoff.get("question"))
        _add_label_paragraph(document, "Why AI stopped", handoff.get("why_ai_cannot_resolve"))
        _add_label_paragraph(document, "Observation", issue.get("observation"))
        for passage_id in issue.get("source_passage_ids", []):
            passage = passages.get(passage_id)
            if passage:
                _add_label_paragraph(document, f"Source {passage_id}", passage.get("text"))
        _add_label_paragraph(document, "Researcher decision", issue.get("decision"))
        _add_label_paragraph(document, "Researcher rationale", issue.get("decision_rationale"))
        _add_label_paragraph(document, "Expert advice", handoff.get("expert_advice"))
        _add_label_paragraph(document, "Expert rationale", handoff.get("expert_rationale"), after=8)
        _add_label_paragraph(document, "Researcher response", handoff.get("researcher_response"))
        _add_label_paragraph(document, "Linked protocol revision", handoff.get("researcher_revised_text"), after=8)

    document.add_heading("Review boundary", level=1)
    document.add_paragraph(
        "This summary records advice and unresolved questions. It is not an institutional approval decision and does not replace the institution's required review record."
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_expert_portfolio_docx(sessions: List[Dict[str, Any]]) -> bytes:
    """Build an expert-facing summary across the protocols they can access."""
    document = Document()
    _configure_docx(document, "SAFEBARS  /  EXPERT CASELOAD SUMMARY")
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    title = document.add_paragraph()
    _set_run(title.add_run("SAFEBARS EXPERT CASELOAD SUMMARY"), size=22, color=INK, bold=True)
    subtitle = document.add_paragraph()
    _set_run(subtitle.add_run("Standard-user ethics applications and recorded expert advice"), size=14, color=MUTED)
    _add_docx_boundary_callout(document)

    total_handoffs = sum(len(session.get("handoffs", [])) for session in sessions)
    unresolved = sum(
        item.get("status") != "resolved"
        for session in sessions
        for item in session.get("handoffs", [])
    )
    advice_given = sum(
        bool(item.get("expert_advice"))
        for session in sessions
        for item in session.get("handoffs", [])
    )
    _add_label_paragraph(document, "Applications in scope", len(sessions))
    _add_label_paragraph(document, "Total handoffs", total_handoffs)
    _add_label_paragraph(document, "Unresolved handoffs", unresolved)
    _add_label_paragraph(document, "Handoffs with recorded advice", advice_given)
    if generated_at:
        _add_label_paragraph(document, "Generated", generated_at)

    document.add_heading("Caseload priorities", level=1)
    prioritized = sorted(
        sessions,
        key=lambda session: (
            -sum(item.get("priority") == "high" and item.get("status") != "resolved" for item in session.get("handoffs", [])),
            -sum(item.get("status") != "resolved" for item in session.get("handoffs", [])),
            _text(session.get("project", {}).get("title")),
        ),
    )
    for index, session in enumerate(prioritized, start=1):
        project = session.get("project", {})
        readiness = session.get("application_readiness") or build_application_readiness(session)
        handoffs = session.get("handoffs", [])
        high_unresolved = sum(item.get("priority") == "high" and item.get("status") != "resolved" for item in handoffs)
        unresolved_count = sum(item.get("status") != "resolved" for item in handoffs)
        document.add_heading(f"{index}. {_text(project.get('title'), 'Untitled research project')}", level=2)
        _add_label_paragraph(document, "Session", session.get("id"))
        _add_label_paragraph(document, "Research context", project.get("context"))
        _add_label_paragraph(document, "Participants and affected relationships", project.get("target_people"))
        _add_label_paragraph(document, "Application profile", readiness.get("profile", {}).get("label"))
        _add_label_paragraph(document, "Application completeness", f"{readiness.get('completion_percent', 0)}% documented")
        _add_label_paragraph(document, "Review priority", f"{high_unresolved} high-priority unresolved; {unresolved_count} unresolved overall")

        deliberations = session.get("tradeoff_deliberations", {})
        if deliberations:
            document.add_heading("Recorded research-design trade-offs", level=3)
            tradeoff_by_id = {
                item.get("id"): item
                for item in session.get("framework_assessment", {}).get("tradeoffs", [])
            }
            for tradeoff_id, deliberation in deliberations.items():
                tradeoff = tradeoff_by_id.get(tradeoff_id, {})
                left = tradeoff.get("left", {})
                right = tradeoff.get("right", {})
                left_value = int(deliberation.get("value", left.get("value", 50)))
                _add_label_paragraph(
                    document,
                    tradeoff.get("title", tradeoff_id),
                    f"{_text(left.get('label'))}: {left_value} / {_text(right.get('label'))}: {100 - left_value}. Rationale: {_text(deliberation.get('rationale'))}",
                )

        document.add_heading("Application gaps", level=3)
        gaps = [field for field in readiness.get("fields", []) if field.get("status") != "documented"]
        if not gaps:
            document.add_paragraph("No application-profile gaps were detected in the submitted fields.")
        for field in gaps:
            _add_label_paragraph(
                document,
                f"{field.get('label', 'Application field')} - {_text(field.get('status')).upper()}",
                field.get("prompt"),
            )

        document.add_heading("Advice and response record", level=3)
        if not handoffs:
            document.add_paragraph("No handoffs or expert advice are recorded for this application.")
        for handoff in sorted(
            handoffs,
            key=lambda item: ({"high": 0, "medium": 1, "standard": 2}.get(item.get("priority", "standard"), 3), item.get("status") == "resolved"),
        ):
            _add_label_paragraph(document, "Question", handoff.get("question"))
            _add_label_paragraph(document, "Priority / status", f"{handoff.get('priority', 'standard')} / {handoff.get('status', 'open')}")
            _add_label_paragraph(document, "Recommended reviewer", handoff.get("recommended_role_label", handoff.get("owner")))
            _add_label_paragraph(document, "Why AI stopped", handoff.get("why_ai_cannot_resolve"))
            _add_label_paragraph(document, "Expert advice", handoff.get("expert_advice"))
            _add_label_paragraph(document, "Expert rationale", handoff.get("expert_rationale"))
            _add_label_paragraph(document, "Researcher response", handoff.get("researcher_response"))
            _add_label_paragraph(document, "Linked protocol revision", handoff.get("researcher_revised_text"), after=8)

    document.add_heading("Review boundary", level=1)
    document.add_paragraph(
        "This caseload summary helps an invited expert prioritize review and preserve advice across accessible applications. "
        "It is not an institution-wide register, an approval decision, or a substitute for the institution's official review record."
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()
